#!/usr/bin/env python3
"""
Incremental PRD Generator - Sections 9-11
API Specifications, Data Requirements, UI/UX Requirements
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class PlamboPRDSections911:
    """Generator for Sections 9-11 of PRD"""
    
    def __init__(self, existing_doc_path):
        """Load existing document"""
        self.doc = Document(existing_doc_path)
        print(f"✓ Loaded existing document: {existing_doc_path}")
    
    def add_heading_1(self, text):
        """Add Heading 1"""
        heading = self.doc.add_heading(text, level=1)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(12)
        heading_format.space_after = Pt(6)
        return heading
    
    def add_heading_2(self, text):
        """Add Heading 2"""
        heading = self.doc.add_heading(text, level=2)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(10)
        heading_format.space_after = Pt(4)
        return heading
    
    def add_heading_3(self, text):
        """Add Heading 3"""
        heading = self.doc.add_heading(text, level=3)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(8)
        heading_format.space_after = Pt(2)
        return heading
    
    def add_paragraph(self, text, bold=False, style=None):
        """Add paragraph"""
        para = self.doc.add_paragraph(text, style=style)
        para_format = para.paragraph_format
        para_format.space_before = Pt(0)
        para_format.space_after = Pt(6)
        para_format.line_spacing = 1.15
        if bold:
            for run in para.runs:
                run.bold = True
        return para
    
    def add_bullet_list(self, items):
        """Add bulleted list"""
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')
            para_format = para.paragraph_format
            para_format.space_after = Pt(3)
            para_format.line_spacing = 1.15
    
    def add_code_block(self, code, language="json"):
        """Add code block"""
        para = self.doc.add_paragraph(code, style='Normal')
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        return para
    
    def add_table(self, rows, cols, headers=None):
        """Add formatted table"""
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'
        
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tcPr = header_cells[i]._element.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:shd')
                tcVAlign.set(qn('w:fill'), '4472C4')
                tcPr.append(tcVAlign)
        
        return table
    
    def generate_section_9_api_specifications(self):
        """Generate Section 9: API Specifications"""
        self.add_heading_1("9. API Specifications")
        
        self.add_paragraph(
            "This section details the RESTful API endpoints, request/response formats, authentication mechanisms, "
            "and error handling for the Plambo Backend. All endpoints follow REST conventions and return JSON responses."
        )
        
        # 9.1 Authentication & Authorization
        self.add_heading_2("9.1 Authentication & Authorization")
        
        self.add_heading_3("9.1.1 JWT Token Format")
        
        self.add_paragraph("Example JWT Token Header:")
        self.add_code_block("""{
  "alg": "HS256",
  "typ": "JWT",
  "kid": "key_version_1"
}""")
        
        self.add_paragraph("JWT Payload Structure:")
        self.add_code_block("""{
  "user_id": "uuid-string",
  "client_id": "profile-uuid",
  "roles": ["user", "analyst"],
  "exp": 1705000000,
  "iat": 1704996400,
  "iss": "plambo-api-v1",
  "sub": "user@organization.com"
}""")
        
        self.add_bullet_list([
            "Token Expiration: 1 hour",
            "Refresh Token: 30 days (issued with initial token)",
            "Algorithm: HS256 with 256-bit key",
            "Key Rotation: Quarterly (new 'kid' issued)",
            "Token Refresh Endpoint: POST /api/auth/refresh"
        ])
        
        # 9.2 Core Query API
        self.add_heading_2("9.2 Core Endpoints")
        
        self.add_heading_3("9.2.1 Query Endpoint - POST /api/query")
        
        self.add_paragraph("Request Body:", bold=True)
        self.add_code_block("""{
  "query_text": "What was the total revenue by region for Q4 2024?",
  "session_id": "550e8400-e29b-41d4-a716-446655440000",
  "client_id": "client-uuid-string",
  "parameters": {
    "max_results": 1000,
    "confidence_threshold": 0.7,
    "retrieval_k": 5,
    "include_reasoning": true
  }
}""")
        
        self.add_paragraph("Success Response (200 OK):", bold=True)
        self.add_code_block("""{
  "status": "success",
  "query_id": "query-uuid",
  "insights": {
    "summary": "Q4 2024 revenue by region shows...",
    "key_findings": [
      "North America led with $2.3M (45%)",
      "EMEA contributed $1.8M (35%)",
      "APAC showed $1.2M (20%)"
    ],
    "confidence": 0.92
  },
  "data": {
    "rows": [
      {"region": "North America", "revenue": 2300000},
      {"region": "EMEA", "revenue": 1800000},
      {"region": "APAC", "revenue": 1200000}
    ],
    "row_count": 3,
    "query_time_ms": 1240
  },
  "metadata": {
    "intent": "AGGREGATE",
    "entities_extracted": ["revenue", "region", "Q4 2024"],
    "documents_retrieved": 3,
    "processing_steps": ["parse", "semanticsearch", "sqlgen", "execute", "synthesize"]
  },
  "traceability": {
    "session_id": "550e8400-e29b-41d4-a716-446655440000",
    "timestamp": "2024-02-13T10:30:45.123Z",
    "user_id": "user-uuid"
  }
}""")
        
        self.add_paragraph("Error Response (400 Bad Request):", bold=True)
        self.add_code_block("""{
  "status": "error",
  "error_code": "INVALID_QUERY",
  "message": "Query text cannot exceed 500 characters",
  "details": {
    "field": "query_text",
    "constraint": "max_length:500",
    "provided": 612
  }
}""")
        
        # 9.2.2 Session Endpoints
        self.add_heading_3("9.2.2 Session Management Endpoints")
        
        session_endpoints = [
            {
                "method": "POST",
                "endpoint": "/api/sessions",
                "description": "Create new session",
                "body": '{"client_id": "uuid"}',
                "response": '{"session_id": "uuid", "created_at": "timestamp"}'
            },
            {
                "method": "GET",
                "endpoint": "/api/sessions/{session_id}",
                "description": "Get session metadata",
                "response": '{"session_id": "uuid", "user_id": "uuid", "query_count": 5, "created_at": "timestamp"}'
            },
            {
                "method": "GET",
                "endpoint": "/api/sessions/{session_id}/history",
                "description": "Get full conversation history",
                "response": '[{"query": "...", "response": {...}, "timestamp": "..."}]'
            },
            {
                "method": "DELETE",
                "endpoint": "/api/sessions/{session_id}",
                "description": "Close/delete session",
                "response": '{"status": "closed", "session_id": "uuid"}'
            }
        ]
        
        endpoints_table = self.add_table(len(session_endpoints) + 1, 5)
        headers = endpoints_table.rows[0].cells
        headers[0].text = "Method"
        headers[1].text = "Endpoint"
        headers[2].text = "Description"
        headers[3].text = "Request/Response"
        headers[4].text = "Status Code"
        
        for idx, ep in enumerate(session_endpoints, 1):
            endpoints_table.rows[idx].cells[0].text = ep["method"]
            endpoints_table.rows[idx].cells[1].text = ep["endpoint"]
            endpoints_table.rows[idx].cells[2].text = ep["description"]
            endpoints_table.rows[idx].cells[3].text = f"Request: {ep.get('body', 'N/A')}\nResponse: {ep['response']}"
            endpoints_table.rows[idx].cells[4].text = "200 / 2xx"
        
        # 9.3 Health & Monitoring
        self.add_heading_2("9.3 Health & Monitoring Endpoints")
        
        self.add_heading_3("9.3.1 Health Check - GET /api/health")
        
        self.add_code_block("""{
  "status": "healthy",
  "timestamp": "2024-02-13T10:35:20.456Z",
  "components": {
    "database": {"status": "healthy", "latency_ms": 2},
    "vectorstore": {"status": "healthy", "latency_ms": 45},
    "llm_api": {"status": "healthy", "latency_ms": 120},
    "cache": {"status": "healthy", "latency_ms": 1},
    "auth_service": {"status": "healthy", "latency_ms": 5}
  },
  "version": "1.0.0",
  "uptime_seconds": 86400
}""")
        
        self.add_heading_3("9.3.2 Readiness Probe - GET /api/health/ready")
        self.add_paragraph("Returns 200 if system ready to accept queries; 503 if initializing/maintenance")
        
        self.add_heading_3("9.3.3 Liveness Probe - GET /api/health/live")
        self.add_paragraph("Returns 200 if system process alive; 503 if deadlocked/crashed")
        
        # 9.4 Rate Limiting & Quotas
        self.add_heading_2("9.4 Rate Limiting & Quotas")
        
        rate_limit_info = [
            "Global Rate Limit: 100 requests/minute per user (1000/minute per IP)",
            "Query Limit: 50 queries/hour per user profile",
            "Session Limit: 50 concurrent sessions per user",
            "Response Headers:",
            "  - X-RateLimit-Limit: Total requests allowed",
            "  - X-RateLimit-Remaining: Remaining requests",
            "  - X-RateLimit-Reset: Unix timestamp of reset time",
            "Burst Handling: Queue excess requests, return 429 Too Many Requests when queue full",
            "Backoff Strategy: Exponential backoff with jitter (recommended for clients)"
        ]
        
        self.add_bullet_list(rate_limit_info)
        
        # 9.5 Request/Response Format Standards
        self.add_heading_2("9.5 Request/Response Standards")
        
        standards = [
            ("Content-Type", "application/json (all requests/responses)"),
            ("Request Size", "Maximum 10MB payload"),
            ("Response Compression", "Gzip for responses >1KB"),
            ("Date Format", "ISO 8601 with UTC timezone (e.g., 2024-02-13T10:35:20.456Z)"),
            ("Numeric Format", "IEEE 754 double-precision floating point"),
            ("Currency", "Numeric values, currency code in separate field"),
            ("Boolean", "true/false (lowercase JSON)"),
            ("Null", "null for missing values (not omitted)"),
            ("Array Order", "Deterministic order (sorted by primary key where applicable)"),
            ("Pagination", "Limit/offset with defaults: limit=100, offset=0")
        ]
        
        standards_table = self.add_table(len(standards) + 1, 2, headers=["Standard", "Specification"])
        for idx, (std, spec) in enumerate(standards, 1):
            standards_table.rows[idx].cells[0].text = std
            standards_table.rows[idx].cells[1].text = spec
        
        self.doc.add_page_break()
    
    def generate_section_10_data_requirements(self):
        """Generate Section 10: Data Requirements & Management"""
        self.add_heading_1("10. Data Requirements & Management")
        
        # 10.1 Input Data Format & Structure
        self.add_heading_2("10.1 Input Data Format & Structure")
        
        self.add_heading_3("10.1.1 Supported File Formats")
        
        formats_data = [
            ("Format", "File Extension", "When Used", "Processing"),
            ("CSV", ".csv", "Tabular data, exports", "Auto-detect schema, type inference"),
            ("Excel", ".xlsx, .xls", "Business reports, data dumps", "Per-sheet processing"),
            ("Parquet", ".parquet", "Production data loads, optimal queries", "Direct FAISS indexing"),
            ("JSON", ".json", "Nested/semi-structured data", "Flattening for tabular analysis"),
            ("PDF", ".pdf", "Documents, reports, FAQs", "Text extraction, then embedding"),
        ]
        
        formats_table = self.add_table(len(formats_data), 4, headers=formats_data[0])
        for idx in range(1, len(formats_data)):
            formats_table.rows[idx].cells[0].text = formats_data[idx][0]
            formats_table.rows[idx].cells[1].text = formats_data[idx][1]
            formats_table.rows[idx].cells[2].text = formats_data[idx][2]
            formats_table.rows[idx].cells[3].text = formats_data[idx][3]
        
        self.add_heading_3("10.1.2 Schema & Data Types")
        
        self.add_paragraph("Supported Column Data Types:")
        
        data_types = [
            ("INTEGER", "32-bit signed integer, range -2,147,483,648 to 2,147,483,647"),
            ("BIGINT", "64-bit signed integer for large counts/IDs"),
            ("DECIMAL(p,s)", "Precision 'p' digits, scale 's' decimals (e.g., DECIMAL(15,2) for currency)"),
            ("FLOAT", "IEEE 754 single-precision (32-bit)"),
            ("DOUBLE", "IEEE 754 double-precision (64-bit)"),
            ("VARCHAR", "Variable-length strings, max 65535 characters per DuckDB"),
            ("BOOLEAN", "True/False values"),
            ("DATE", "ISO 8601 format YYYY-MM-DD"),
            ("TIMESTAMP", "ISO 8601 with timezone YYYY-MM-DDTHH:MM:SS.sssZ"),
            ("NULL", "Missing/unknown values in any column")
        ]
        
        for dtype, desc in data_types:
            self.add_paragraph(f"{dtype}: {desc}")
        
        self.add_heading_3("10.1.3 Data Quality Requirements")
        
        quality_reqs = [
            "Completeness: Minimum 95% non-null values per column",
            "Consistency: Referential integrity between related tables",
            "Accuracy: Numeric values ±0.01% of source system",
            "Timeliness: Not older than 24 hours for queries",
            "Uniqueness: Primary key violations detected and reported",
            "Valid Ranges: Column values within business-defined ranges",
            "Format Conformance: Dates/timestamps in ISO 8601 format"
        ]
        
        self.add_bullet_list(quality_reqs)
        
        # 10.2 Knowledge Base Management
        self.add_heading_2("10.2 Knowledge Base Management")
        
        self.add_heading_3("10.2.1 Knowledge Base Purpose & Content")
        
        self.add_paragraph(
            "Knowledge Base serves as domain-specific context for RAG pipeline. It may include FAQs, business rules, "
            "product documentation, and domain-specific terminology."
        )
        
        kb_content = [
            "Business Rules: 'Revenue recognition occurs on invoice date, not payment date'",
            "Glossary: 'ASP = Average Selling Price; calculated as total revenue / unit count'",
            "Product Definitions: 'Product line X discontinued Q2 2024'",
            "Process Documentation: Step-by-step business processes",
            "FAQ Documents: Common questions and standardized answers",
            "Historical Context: Significant business events (mergers, expansions)",
            "Metric Definitions: 'Margin = (Revenue - COGS) / Revenue * 100'"
        ]
        
        self.add_bullet_list(kb_content)
        
        self.add_heading_3("10.2.2 Document Upload Process")
        
        upload_flow = [
            "1. User uploads file via POST /api/knowledge-base/upload",
            "2. File validated (size, format, content scanning for malware)",
            "3. Document extracted to text content",
            "4. Text chunked into ~500-token segments",
            "5. Each chunk embedded using sentence-transformers model",
            "6. Embeddings indexed in FAISS under profile-specific index",
            "7. Document metadata stored in PostgreSQL",
            "8. User notified of completion (email or webhook)"
        ]
        
        self.add_bullet_list(upload_flow)
        
        # 10.3 Session Data Persistence
        self.add_heading_2("10.3 Session Data Persistence")
        
        self.add_heading_3("10.3.1 Session Data Model")
        
        self.add_code_block("""CREATE TABLE sessions (
  session_id UUID PRIMARY KEY,
  user_id UUID NOT NULL,
  client_id UUID NOT NULL,
  created_at TIMESTAMP NOT NULL,
  last_activity TIMESTAMP NOT NULL,
  status TEXT CHECK (status IN ('active', 'idle', 'closed', 'archived')),
  conversation_history JSONB NOT NULL,
  context_summary TEXT,
  query_count INTEGER DEFAULT 0,
  tokens_used INTEGER DEFAULT 0,
  FOREIGN KEY (user_id) REFERENCES users(user_id),
  FOREIGN KEY (client_id) REFERENCES profiles(profile_id)
);

CREATE TABLE conversation_exchanges (
  exchange_id UUID PRIMARY KEY,
  session_id UUID NOT NULL REFERENCES sessions(session_id),
  sequence_number INTEGER NOT NULL,
  user_query TEXT NOT NULL,
  system_response JSONB NOT NULL,
  intent TEXT,
  confidence_score FLOAT,
  timestamp TIMESTAMP NOT NULL,
  UNIQUE (session_id, sequence_number)
);""")
        
        self.add_heading_3("10.3.2 Data Retention Policy")
        
        retention_policy = [
            "Active Sessions: Kept in PostgreSQL + Redis cache",
            "Idle Sessions (30-90 days): Moved to snapshot storage (S3)",
            "Archived Sessions (>90 days): Long-term storage with index",
            "Retention Duration: 30 years for audit/compliance",
            "Deletion: Upon explicit user request or data subject right",
            "Backup: Daily full backups, hourly incremental (30-day retention)",
            "Disaster Recovery: 24-hour RPO, 15-minute RTO"
        ]
        
        self.add_bullet_list(retention_policy)
        
        # 10.4 Data Privacy & PII Handling
        self.add_heading_2("10.4 Data Privacy & PII Handling")
        
        self.add_heading_3("10.4.1 PII Detection & Masking")
        
        pii_handling = [
            "Automatic Detection: PII patterns (SSN, email, phone, credit card) automatically flagged",
            "Masking in Logs: PII values replaced with [REDACTED] in all logs",
            "Masking in Responses: Optional PII masking endpoint for query results",
            "Encryption: Sensitive columns encrypted with AES-256",
            "Access Control: PII columns require specific permission role",
            "Audit Logging: All PII access logged with user/timestamp/reason",
            "Anonymization: Support for tokenization/hashing for analytics"
        ]
        
        self.add_bullet_list(pii_handling)
        
        self.add_heading_3("10.4.2 GDPR Compliance")
        
        gdpr_items = [
            "Right to be Forgotten: DELETE /api/users/{user_id} cascades to all data",
            "Data Portability: Export all user data in structured format (JSON/CSV)",
            "Consent Tracking: Store explicit consent with timestamps",
            "Data Processing Agreement: Pre-signed DPA with organization",
            "Breach Notification: Alert within 72 hours per Article 33",
            "Privacy by Design: Data minimization, purpose limitation"
        ]
        
        self.add_bullet_list(gdpr_items)
        
        self.doc.add_page_break()
    
    def generate_section_11_uiux_requirements(self):
        """Generate Section 11: UI/UX Requirements"""
        self.add_heading_1("11. UI/UX Requirements")
        
        self.add_paragraph(
            "While Plambo Backend is a service/API platform, the UX requirements define how the system should communicate "
            "with clients, how responses should be formatted for optimal consumption, and accessibility considerations."
        )
        
        # 11.1 API Response UX
        self.add_heading_2("11.1 API Response User Experience")
        
        self.add_heading_3("11.1.1 Error Message Quality")
        
        error_ux = [
            "Clarity: Non-technical users should understand error without documentation",
            "Actionability: Include suggestion on how to resolve (e.g., 'Check query syntax' or 'Retry in 60 seconds')",
            "Context: Include field/parameter that caused error, not just error code",
            "Localization: Error messages available in all supported languages",
            "Examples:",
            "  • BAD: 'Invalid input'",
            "  • GOOD: 'Query contains unknown column \"revenue_xyz\". Did you mean \"total_revenue\"?'"
        ]
        
        self.add_bullet_list(error_ux)
        
        self.add_heading_3("11.1.2 Cognitive Load Reduction")
        
        cognitive_load = [
            "Response Summarization: Insights <200 chars for quick scanning",
            "Hierarchical Information: Summary first, then details on demand",
            "Visual Cues: Status emoji (✓ success, ⚠ warning) in text responses",
            "Numeric Formatting: Millions shown as '2.3M', not '2300000'",
            "Percentage Formatting: 0.92 shown as '92%' in insights",
            "Confidence Scoring: Always include 0-100% confidence in insights",
            "Processing Time: Include query_time_ms for transparency"
        ]
        
        self.add_bullet_list(cognitive_load)
        
        # 11.2 Documentation & Discoverability
        self.add_heading_2("11.2 Documentation & API Discoverability")
        
        self.add_heading_3("11.2.1 OpenAPI Specification Completeness")
        
        openapi_items = [
            "Operation ID: Every endpoint has unique operationId (e.g., 'queryAnalyticalQuestion')",
            "Descriptions: Every operation has >50 char human-readable description",
            "Example Values: Request/response examples in OpenAPI spec",
            "Success & Error Examples: Show both 200 and error responses",
            "Parameter Documentation: Description for every query/path parameter",
            "Schema Descriptions: Every property in models documented",
            "Deprecated Marking: Mark deprecated endpoints with x-deprecated flag"
        ]
        
        self.add_bullet_list(openapi_items)
        
        self.add_heading_3("11.2.2 Interactive Documentation")
        
        self.add_paragraph("Swagger UI Features:")
        
        swagger_features = [
            "Try-it-out: Execute sample requests directly from documentation",
            "Authentication: OAuth2/JWT token input field for auth testing",
            "Response Display: Pretty-printed JSON with syntax highlighting",
            "Schema Explorer: Browse model definitions with descriptions",
            "Search: Full-text search across endpoints and documentation",
            "Code Samples: Auto-generated code examples (Python, JavaScript, cURL, Java)"
        ]
        
        self.add_bullet_list(swagger_features)
        
        # 11.3 Response Format Consistency
        self.add_heading_2("11.3 Response Format Consistency")
        
        self.add_heading_3("11.3.1 Response Envelope Standard")
        
        self.add_paragraph("All responses follow envelope structure:")
        
        self.add_code_block("""{
  "status": "success" | "error" | "warning",
  "data": {...},           # Actual response data
  "metadata": {...},       # Processing metadata
  "errors": [...],         # Array of error objects (if status=error)
  "timestamp": "2024-02-13T10:35:20.456Z",
  "request_id": "req-uuid" # For tracing
}""")
        
        self.add_heading_3("11.3.2 Consistent Field Naming")
        
        naming_conventions = [
            "camelCase for JSON field names (e.g., 'totalRevenue', not 'total_revenue')",
            "Consistent pluralization: 'items' (array), 'count' (number)",
            "Timestamp field name: Always 'timestamp' in ISO 8601 format",
            "Unique identifiers: Always suffixed with '_id' (e.g., 'session_id', 'user_id')",
            "Boolean fields: Prefixed with 'is_' or 'has_' (e.g., 'isActive', 'hasError')",
            "Money fields: Always with currency code in separate field",
            "Percentage fields: Numeric 0-100, not 0-1 decimal"
        ]
        
        self.add_bullet_list(naming_conventions)
        
        # 11.4 Accessibility & Usability
        self.add_heading_2("11.4 Accessibility & Usability")
        
        self.add_heading_3("11.4.1 API Accessibility")
        
        accessibility_items = [
            "Query Language: Support natural language in multiple languages/dialects",
            "Error Recovery: Provide suggestions for correcting malformed queries",
            "Partial Success: Return partial results when some dimensions unavailable",
            "Graceful Degradation: Fallback to simpler analysis if complex calculation fails",
            "Offline Capability: Cache frequently-used results for offline access (future)",
            "Mobile-Friendly: Compact response format option (minimized field names)"
        ]
        
        self.add_bullet_list(accessibility_items)
        
        self.add_heading_3("11.4.2 Performance Feedback")
        
        perf_feedback = [
            "Progress Indication: For long-running queries, return 202 Accepted with status endpoint",
            "Estimated Time: Include estimated_seconds_remaining for queries >2 seconds",
            "Cancellation Support: Allow cancellation via DELETE /api/query/{query_id}",
            "Latency Transparency: Include component latencies (parse: Xms, retrieve: Yms, generate: Zms)",
            "Query Cost: Optional cost estimation for monetized platforms"
        ]
        
        self.add_bullet_list(perf_feedback)
        
        # 11.5 Multi-Tenancy & Personalization
        self.add_heading_2("11.5 Multi-Tenancy & Personalization UX")
        
        self.add_heading_3("11.5.1 Profile-Specific Customization")
        
        customization = [
            "Configuration Exposure: Client-specific settings reflected in API responses",
            "Terminology: Dimension names can be customized per profile (default -> 'Sales Region')",
            "Number Formatting: Profile-specific currency, decimal separator, thousand separator",
            "Timezone Handling: Dates/times adjusted to profile's timezone in responses",
            "Language Preference: Response language follows user's language_preference setting"
        ]
        
        self.add_bullet_list(customization)
        
        self.add_heading_3("11.5.2 Personalized Recommendations")
        
        recommendations = [
            "Suggested Queries: Based on user's query history and profile",
            "Similar Sessions: Link to previously similar analyses",
            "Related Insights: 'You might also be interested in...' suggestions",
            "Trending Metrics: 'Trending this week: YoY revenue growth +12%'",
            "Anomaly Alerts: Proactive notification of significant variations"
        ]
        
        self.add_bullet_list(recommendations)
        
        # 11.6 Accessibility Standards
        self.add_heading_2("11.6 WCAG 2.1 Accessibility Considerations")
        
        self.add_heading_3("11.6.1 For Frontend Integrations")
        
        wcag_items = [
            "WCAG 2.1 Level AA: Frontend implementations should meet AA compliance minimum",
            "Color Contrast: Text contrast ratio ≥4.5:1 for normal text",
            "Keyboard Navigation: All functionality accessible via keyboard",
            "Screen Reader Support: Proper ARIA labels and semantic HTML",
            "Focus Management: Clear focus indicators for all interactive elements",
            "Alt Text: All charts/visuals include text descriptions",
            "Plain Language: Documentation uses simple language, <15th grade reading level"
        ]
        
        self.add_bullet_list(wcag_items)
        
        self.add_heading_3("11.6.2 API Documentation Accessibility")
        
        api_accessibility = [
            "Documentation in Multiple Formats: HTML, Markdown, PDF",
            "Code Examples: Syntax-highlighted with language identification",
            "Keyboard-Accessible Documentation: All links/controls keyboard navigable",
            "Structured Headings: Proper heading hierarchy (H1, H2, H3)",
            "List Formatting: Bullet lists for parallel structure",
            "Link Text: Descriptive link text ('See API documentation' not 'Click here')",
            "Table Headers: Proper <th> elements for screen reader identification"
        ]
        
        self.add_bullet_list(api_accessibility)
        
        self.doc.add_page_break()
    
    def save_document(self, filename):
        """Save document"""
        self.doc.save(filename)
        print(f"✓ Document saved: {filename}")
        return filename


def main():
    """Generate Sections 9-11"""
    print("\n" + "=" * 70)
    print("PLAMBO PRD GENERATOR - SECTIONS 9-11 EXTENSION")
    print("=" * 70)
    
    existing_file = r"c:\9Trix\CODE\Plambo\plamboBE\Plambo_PRD_v1.0_Sections_1-8.docx"
    
    print(f"\n[STEP 1] Loading existing document...")
    generator = PlamboPRDSections911(existing_file)
    
    print("[STEP 2] Generating Section 9: API Specifications...")
    generator.generate_section_9_api_specifications()
    
    print("[STEP 3] Generating Section 10: Data Requirements...")
    generator.generate_section_10_data_requirements()
    
    print("[STEP 4] Generating Section 11: UI/UX Requirements...")
    generator.generate_section_11_uiux_requirements()
    
    output_file = r"c:\9Trix\CODE\Plambo\plamboBE\Plambo_PRD_v1.0_Sections_1-11.docx"
    generator.save_document(output_file)
    
    print("\n" + "=" * 70)
    print("✓ SECTIONS 9-11 COMPLETE - DOCUMENT EXTENDED")
    print("=" * 70)
    print(f"\nOutput: Plambo_PRD_v1.0_Sections_1-11.docx")
    print("\nSections now included:")
    print("  ✓ Section 1-8: Previously generated")
    print("  ✓ Section 9: API Specifications")
    print("  ✓ Section 10: Data Requirements & Management")
    print("  ✓ Section 11: UI/UX Requirements")
    print("\n→ Ready for final sections. Type 'Continue' to proceed to Sections 12-16.")
    
    return output_file


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""
Incremental PRD Generator - Extension Module
Sections 3-4: User Research & High-Level Requirements
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class PlamboPRDExtension:
    """Extension for adding sections to existing PRD document"""
    
    def __init__(self, existing_doc_path):
        """Load existing document"""
        self.doc = Document(existing_doc_path)
        print(f"✓ Loaded existing document: {existing_doc_path}")
    
    def add_heading_1(self, text):
        """Add Heading 1 with formatting"""
        heading = self.doc.add_heading(text, level=1)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(12)
        heading_format.space_after = Pt(6)
        return heading
    
    def add_heading_2(self, text):
        """Add Heading 2 with formatting"""
        heading = self.doc.add_heading(text, level=2)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(10)
        heading_format.space_after = Pt(4)
        return heading
    
    def add_heading_3(self, text):
        """Add Heading 3 with formatting"""
        heading = self.doc.add_heading(text, level=3)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(8)
        heading_format.space_after = Pt(2)
        return heading
    
    def add_paragraph(self, text, bold=False, style=None):
        """Add paragraph with optional formatting"""
        para = self.doc.add_paragraph(text, style=style)
        para_format = para.paragraph_format
        para_format.space_before = Pt(0)
        para_format.space_after = Pt(6)
        para_format.line_spacing = 1.15
        if bold:
            for run in para.runs:
                run.bold = True
        return para
    
    def add_bullet_list(self, items):
        """Add bulleted list"""
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')
            para_format = para.paragraph_format
            para_format.space_after = Pt(3)
            para_format.line_spacing = 1.15
    
    def add_numbered_list(self, items):
        """Add numbered list"""
        for item in items:
            para = self.doc.add_paragraph(item, style='List Number')
            para_format = para.paragraph_format
            para_format.space_after = Pt(3)
            para_format.line_spacing = 1.15
    
    def add_table(self, rows, cols, headers=None, data=None):
        """Add formatted table"""
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'
        
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set header background color
                tcPr = header_cells[i]._element.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:shd')
                tcVAlign.set(qn('w:fill'), '4472C4')
                tcPr.append(tcVAlign)
        
        if data:
            for row_idx, row_data in enumerate(data, 1):
                for col_idx, cell_data in enumerate(row_data):
                    table.rows[row_idx].cells[col_idx].text = str(cell_data)
        
        return table
    
    def generate_section_3_user_research(self):
        """Generate Section 3: User Research & Analysis"""
        self.add_heading_1("3. User Research & Analysis")
        
        # 3.1 User Environment
        self.add_heading_2("3.1 User Environment")
        
        self.add_heading_3("3.1.1 Deployment Context")
        self.add_paragraph(
            "Plambo is deployed in enterprise environments where users access the platform through web browsers, "
            "mobile applications (future), and programmatic integrations via APIs. The platform operates within corporate "
            "networks with managed security policies, internal SSO, and compliance requirements."
        )
        
        self.add_heading_3("3.1.2 Technical Environment")
        
        tech_env_items = [
            "ER-1: Enterprise firewalls with restricted outbound connectivity",
            "ER-2: VPN/proxy requirements for external API calls",
            "ER-3: On-premises data residency requirements in certain regions",
            "ER-4: Integration with existing data warehouses (Snowflake, BigQuery, Redshift)",
            "ER-5: LDAP/Active Directory SSO for authentication",
            "ER-6: Compliance frameworks: HIPAA, GDPR, SOC 2, FedRAMP"
        ]
        self.add_bullet_list(tech_env_items)
        
        self.add_heading_3("3.1.3 Usage Patterns")
        
        usage_patterns = [
            "UP-1: Peak usage during business hours (9 AM - 5 PM); off-peak at nights/weekends",
            "UP-2: Ad-hoc queries dominate (70%); scheduled/recurring queries (30%)",
            "UP-3: Average session duration 15-45 minutes",
            "UP-4: Multi-turn conversations with 3-10 query exchanges per session",
            "UP-5: Mobile usage growing at 25% YoY; currently 15% of total traffic"
        ]
        self.add_bullet_list(usage_patterns)
        
        # 3.2 User Problems & Needs Analysis
        self.add_heading_2("3.2 User Problems & Needs Analysis")
        
        self.add_heading_3("3.2.1 Critical Problems (Priority: High)")
        
        problems_data = [
            ("Problem", "User Impact", "Current Solution", "Desired Solution"),
            ("Query Turnaround", "Decisions delayed 1-7 days waiting for analyst", "Manual report generation", "Self-serve access in seconds"),
            ("Technical Skill Gap", "Only SQL-capable users can access data independently", "Bottleneck with analyst team", "Natural language interface"),
            ("Ad-hoc Query Limitations", "Cannot explore data variations easily", "Multiple email requests to analysts", "Instant what-if scenarios"),
            ("Context Loss", "Information scattered across reports and emails", "Manual consolidation, error-prone", "Conversation history and summaries"),
            ("Access Control Complexity", "One-size-fits-all data access", "Security concerns with broad access", "Row-level and column-level access control"),
        ]
        
        problems_table = self.add_table(len(problems_data), 4, headers=problems_data[0])
        for row_idx in range(1, len(problems_data)):
            for col_idx in range(4):
                problems_table.rows[row_idx].cells[col_idx].text = problems_data[row_idx][col_idx]
        
        self.add_heading_3("3.2.2 Secondary Problems (Priority: Medium)")
        
        secondary_problems = [
            "SP-1: Data completeness issues - Missing historical context in queries",
            "SP-2: Analytical accuracy - Insights often require expert validation",
            "SP-3: Mobile limitations - Cannot access insights while traveling",
            "SP-4: Integration friction - Multiple tools required for complete analysis",
            "SP-5: Cost visibility - Hidden spend on analyst time and tools"
        ]
        self.add_bullet_list(secondary_problems)
        
        self.add_heading_3("3.2.3 User Needs Mapping")
        
        self.add_paragraph("Functional Needs:")
        functional_needs = [
            "FN-1: Express complex analytical questions in natural language",
            "FN-2: Receive instant, accurate, contextualized answers",
            "FN-3: Follow-up queries referencing previous context",
            "FN-4: Visual representations alongside text summaries",
            "FN-5: Export results in multiple formats (CSV, PDF, Excel)",
            "FN-6: Drill-down capabilities for detailed analysis",
            "FN-7: Set alerts when metrics cross thresholds",
            "FN-8: Save favorite queries for reuse"
        ]
        self.add_bullet_list(functional_needs)
        
        self.add_paragraph("Non-Functional Needs:")
        nf_needs = [
            "NFN-1: Sub-2-second response times for queries",
            "NFN-2: 99.95% platform availability",
            "NFN-3: Mobile-first responsive interface",
            "NFN-4: Multi-language support (EN, ES, DE, FR, ZH)",
            "NFN-5: Single-sign-on (SSO) integration",
            "NFN-6: Audit logging for compliance",
            "NFN-7: Data encryption at-rest and in-transit",
            "NFN-8: Offline capability (future)"
        ]
        self.add_bullet_list(nf_needs)
        
        # 3.3 Competitive Analysis
        self.add_heading_2("3.3 Competitive Analysis")
        
        self.add_heading_3("3.3.1 Direct Competitors")
        
        competitors_data = [
            ("Product", "Strengths", "Weaknesses", "Plambo Advantage"),
            ("Tableau Pulse", "Excellent visualizations, broad adoption", "Limited NLP, expensive, complex setup", "Purpose-built conversational interface"),
            ("Power BI Q&A", "Microsoft integration, advanced visuals", "Narrow LLM capability, accuracy issues", "Superior semantic search with FAISS"),
            ("Sisense Fusion", "Multi-source integration", "Requires custom configuration, steep learning curve", "Zero-configuration for standard workflows"),
            ("Looker Explore", "SQL-first approach, powerful dashboards", "Not conversational, technical friction", "Pure conversational paradigm"),
        ]
        
        comp_table = self.add_table(len(competitors_data), 4, headers=competitors_data[0])
        for row_idx in range(1, len(competitors_data)):
            for col_idx in range(4):
                comp_table.rows[row_idx].cells[col_idx].text = competitors_data[row_idx][col_idx]
        
        self.add_heading_3("3.3.2 Competitive Differentiation")
        
        differentiation = [
            "• Pure Conversational Interface: Unlike query builders, Plambo is 100% conversational",
            "• Vector-Based Semantic Search: FAISS-powered similarity search exceeds traditional keyword matching",
            "• Multi-Client Isolation: Enterprise-grade tenancy out-of-the-box",
            "• Enterprise-Ready APIs: RESTful architecture for seamless integration",
            "• Open Architecture: Flexible backend enabling custom workflows",
            "• Cost-Effective: Leveraging open-source (DuckDB, FAISS) and LLM APIs",
            "• Speed to Insight: Average 3-5 second response vs competitors' 10-30 seconds"
        ]
        self.add_bullet_list(differentiation)
        
        self.doc.add_page_break()
    
    def generate_section_4_high_level_requirements(self):
        """Generate Section 4: High-Level Product Requirements"""
        self.add_heading_1("4. High-Level Product Requirements")
        
        # 4.1 Feature Summary
        self.add_heading_2("4.1 Feature Summary")
        
        self.add_paragraph(
            "The Plambo Backend comprises six core feature modules that work in concert to deliver "
            "conversational business intelligence. Each module is designed with modularity and extensibility in mind, "
            "enabling future enhancements without architectural disruption."
        )
        
        self.add_heading_3("4.1.1 Core Feature Modules")
        
        # Feature 1
        self.add_paragraph("Feature 1: Conversational Query Engine", bold=True)
        f1_desc = [
            "Description: Natural language query processing that converts user questions into structured analytical operations",
            "Capability: Understand complex, multi-turn conversations with context preservation",
            "Components: Query parser, intent classifier, entity extractor, query optimizer",
            "Expected Accuracy: 85%+ for known question patterns, 70%+ for novel queries"
        ]
        self.add_bullet_list(f1_desc)
        
        # Feature 2
        self.add_paragraph("Feature 2: Multi-Client Profile Management", bold=True)
        f2_desc = [
            "Description: Complete isolation and customization per organizational client",
            "Capability: Independent configurations, knowledge bases, vector stores, and user hierarchies",
            "Components: Profile service, tenant manager, access controller",
            "Max Clients: 50 simultaneously active clients at launch"
        ]
        self.add_bullet_list(f2_desc)
        
        # Feature 3
        self.add_paragraph("Feature 3: Semantic Search & RAG Pipeline", bold=True)
        f3_desc = [
            "Description: Vector-based document retrieval enabling contextual, nuanced insights",
            "Capability: Retrieve top-K relevant documents, synthesize insights via LLM",
            "Components: FAISS indexer, similarity searcher, RAG orchestrator",
            "Performance: Sub-100ms retrieval, configurable K (1-500 documents)"
        ]
        self.add_bullet_list(f3_desc)
        
        # Feature 4
        self.add_paragraph("Feature 4: Session Management & Conversation History", bold=True)
        f4_desc = [
            "Description: Stateful session tracking with full conversation preservation",
            "Capability: Multi-turn conversations with context transfer between queries",
            "Components: Session store, memory buffer, conversation serializer",
            "Retention: Up to 100,000 sessions per client; 30-year conversation archival"
        ]
        self.add_bullet_list(f4_desc)
        
        # Feature 5
        self.add_paragraph("Feature 5: Data Analysis & Insight Generation", bold=True)
        f5_desc = [
            "Description: Advanced analytical capabilities powered by DuckDB and LLM",
            "Capability: Compute aggregations, trends, anomalies, comparisons, forecasts",
            "Components: SQL generator, analytical engine, insight synthesizer",
            "Supported Analyses: Trend analysis, YoY/MoM comparisons, top-N selection, filtering"
        ]
        self.add_bullet_list(f5_desc)
        
        # Feature 6
        self.add_paragraph("Feature 6: File Upload & Knowledge Base Management", bold=True)
        f6_desc = [
            "Description: Asynchronous file processing and vector indexing pipeline",
            "Capability: Accept multiple file formats, transform to parquet, compute embeddings",
            "Components: File handler, transformer pipeline, embedding engine",
            "Supported Formats: CSV, Excel, Parquet, JSON, PDF; Output: Parquet + FAISS indices"
        ]
        self.add_bullet_list(f6_desc)
        
        self.doc.add_page_break()
        
        # 4.2 Feature Prioritization Matrix
        self.add_heading_2("4.2 Feature Prioritization & Roadmap")
        
        self.add_heading_3("4.2.1 MoSCoW Prioritization")
        
        self.add_paragraph("MUST HAVE (MVP - Launch Requirement):", bold=True)
        must_have = [
            "FP-1: Basic conversational query processing (95% accuracy on templated queries)",
            "FP-2: Single-client profile with sample knowledge base",
            "FP-3: FAISS vector-based document retrieval",
            "FP-4: Session management with 5-turn conversation support",
            "FP-5: DuckDB analytical query execution",
            "FP-6: REST API with 3 core endpoints: /query, /sessions, /health",
            "FP-7: JWT authentication and basic authorization",
            "FP-8: Error handling and retry logic"
        ]
        self.add_bullet_list(must_have)
        
        self.add_paragraph("SHOULD HAVE (Q1 Enhancement):", bold=True)
        should_have = [
            "FP-9: Multi-client profile support (up to 10 clients)",
            "FP-10: File upload pipeline with parquet transformation",
            "FP-11: Advanced conversation history with context inference",
            "FP-12: Insight synthesis with trend detection",
            "FP-13: Web search integration for external context",
            "FP-14: Basic audit logging for compliance",
            "FP-15: System monitoring and health metrics",
            "FP-16: Rate limiting and usage analytics"
        ]
        self.add_bullet_list(should_have)
        
        self.add_paragraph("COULD HAVE (Q2 Expansion):", bold=True)
        could_have = [
            "FP-17: Multi-language query support (Spanish, French, German, Mandarin)",
            "FP-18: Advanced forecasting models (ARIMA, Prophet)",
            "FP-19: Real-time alert mechanism for metric thresholds",
            "FP-20: Saved queries/favorites functionality",
            "FP-21: Advanced access control (row-level, column-level security)",
            "FP-22: Webhook support for async operations",
            "FP-23: Export functionality (CSV, Excel, PDF)",
            "FP-24: Mobile-optimized API responses"
        ]
        self.add_bullet_list(could_have)
        
        self.add_paragraph("WON'T HAVE (Post-v2.0):", bold=True)
        wont_have = [
            "FP-25: Custom model training (ChatGPT fine-tuning discontinued)",
            "FP-26: Real-time streaming analytics",
            "FP-27: Off-line capability",
            "FP-28: Quantum computing integration",
            "FP-29: Custom hardware requirement support"
        ]
        self.add_bullet_list(wont_have)
        
        self.add_heading_3("4.2.2 Release Timeline")
        
        timeline_data = [
            ("Release", "Target Date", "Key Features", "Status"),
            ("v1.0 MVP", "Q3 2024", "Core conversational engine, single client, basic APIs", "Planned"),
            ("v1.1 Multi-Client", "Q4 2024", "Multi-client support, file upload, audit logging", "Planned"),
            ("v1.2 Analytics", "Q1 2025", "Advanced insights, web search, forecasting", "Planned"),
            ("v2.0 Enterprise", "Q2 2025", "Multi-language, RBAC, advanced security", "Planned"),
        ]
        
        timeline_table = self.add_table(len(timeline_data), 4, headers=timeline_data[0])
        for row_idx in range(1, len(timeline_data)):
            for col_idx in range(4):
                timeline_table.rows[row_idx].cells[col_idx].text = timeline_data[row_idx][col_idx]
        
        self.doc.add_page_break()
    
    def save_document(self, filename):
        """Save extended document"""
        self.doc.save(filename)
        print(f"✓ Document saved: {filename}")
        return filename


def main():
    """Extend document with Sections 3-4"""
    print("\n" + "=" * 70)
    print("PLAMBO PRD GENERATOR - SECTIONS 3-4 EXTENSION")
    print("=" * 70)
    
    existing_file = r"c:\9Trix\CODE\Plambo\plamboBE\Plambo_PRD_v1.0_Sections_1-2.docx"
    
    print(f"\n[STEP 1] Loading existing document...")
    extension = PlamboPRDExtension(existing_file)
    
    print("[STEP 2] Generating Section 3: User Research & Analysis...")
    extension.generate_section_3_user_research()
    
    print("[STEP 3] Generating Section 4: High-Level Requirements...")
    extension.generate_section_4_high_level_requirements()
    
    output_file = r"c:\9Trix\CODE\Plambo\plamboBE\Plambo_PRD_v1.0_Sections_1-4.docx"
    extension.save_document(output_file)
    
    print("\n" + "=" * 70)
    print("✓ SECTIONS 3-4 COMPLETE - DOCUMENT EXTENDED")
    print("=" * 70)
    print(f"\nOutput: Plambo_PRD_v1.0_Sections_1-4.docx")
    print("\nSections now included:")
    print("  ✓ Section 1: Introduction")
    print("  ✓ Section 2: Product Overview")
    print("  ✓ Section 3: User Research & Analysis")
    print("  ✓ Section 4: High-Level Product Requirements")
    print("\n→ Ready for next sections. Type 'Continue' to proceed to Sections 5-6.")
    
    return output_file


if __name__ == "__main__":
    main()

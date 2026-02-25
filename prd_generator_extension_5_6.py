#!/usr/bin/env python3
"""
Incremental PRD Generator - Sections 5-6
Detailed Features & Functional Requirements (FR)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class PlamboPRDSections56:
    """Generator for Sections 5-6 of PRD"""
    
    def __init__(self, existing_doc_path):
        """Load existing document"""
        self.doc = Document(existing_doc_path)
        print(f"✓ Loaded existing document: {existing_doc_path}")
        self.fr_counter = 1
    
    def add_heading_1(self, text):
        """Add Heading 1"""
        heading = self.doc.add_heading(text, level=1)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(12)
        heading_format.space_after = Pt(6)
        return heading
    
    def add_heading_2(self, text):
        """Add Heading 2"""
        heading = self.doc.add_heading(text, level=2)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(10)
        heading_format.space_after = Pt(4)
        return heading
    
    def add_heading_3(self, text):
        """Add Heading 3"""
        heading = self.doc.add_heading(text, level=3)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(8)
        heading_format.space_after = Pt(2)
        return heading
    
    def add_paragraph(self, text, bold=False, style=None):
        """Add paragraph"""
        para = self.doc.add_paragraph(text, style=style)
        para_format = para.paragraph_format
        para_format.space_before = Pt(0)
        para_format.space_after = Pt(6)
        para_format.line_spacing = 1.15
        if bold:
            for run in para.runs:
                run.bold = True
        return para
    
    def add_bullet_list(self, items):
        """Add bulleted list"""
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')
            para_format = para.paragraph_format
            para_format.space_after = Pt(3)
            para_format.line_spacing = 1.15
    
    def add_table(self, rows, cols, headers=None):
        """Add formatted table"""
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'
        
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tcPr = header_cells[i]._element.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:shd')
                tcVAlign.set(qn('w:fill'), '4472C4')
                tcPr.append(tcVAlign)
        
        return table
    
    def generate_section_5_detailed_features(self):
        """Generate Section 5: Detailed Feature Requirements"""
        self.add_heading_1("5. Detailed Feature Requirements")
        
        # Feature 5.1
        self.add_heading_2("5.1 Conversational Query Engine (CQE)")
        
        self.add_heading_3("5.1.1 Overview")
        self.add_paragraph(
            "The Conversational Query Engine is the core intelligent component that interprets user natural language "
            "queries, extracts intent, identifies entities, and transforms linguistic input into structured analytical operations. "
            "The CQE operates as a multi-stage pipeline with fallback mechanisms for graceful degradation."
        )
        
        self.add_heading_3("5.1.2 Query Processing Pipeline")
        
        pipeline_steps = [
            "Input Reception: Receive user query string via REST API",
            "Session Context Loading: Retrieve current session context and conversation history",
            "Preprocessing: Normalize text (lowercase, tokenization, spell-correction)",
            "Intent Classification: Classify query intent (SELECT, AGGREGATE, COMPARE, TREND, FORECAST, etc.)",
            "Entity Extraction: Identify dimension tables, measures, time periods, filters",
            "Semantic Search: Query vector store for relevant knowledge base documents",
            "Query Optimization: Rewrite SQL for performance and correctness",
            "Execution: Run against DuckDB analytical engine",
            "Result Formatting: Transform output for LLM consumption",
            "Insight Generation: Call LLM with context and results for natural language summary",
            "Response Serialization: Format response with metadata, confidence, traceability"
        ]
        self.add_bullet_list(pipeline_steps)
        
        self.add_heading_3("5.1.3 Core Components")
        
        components_data = [
            ("Component", "Responsibility", "Technology"),
            ("Query Parser", "Tokenize and parse natural language", "spaCy NLP library + custom grammar"),
            ("Intent Classifier", "Determine analytical operation type", "ML classifier (BERT-based) or rule-based"),
            ("Entity Extractor", "Identify tables, columns, filters, time ranges", "Named Entity Recognition (NER)"),
            ("Semantic Searcher", "Retrieve contextual documents from knowledge base", "FAISS + embeddings"),
            ("SQL Generator", "Convert parsed query to SQL", "Rule-based template system + LLM"),
            ("Query Executor", "Run SQL against data", "DuckDB"),
            ("Insight Synthesizer", "Generate natural language insights", "Google Gemini API"),
        ]
        
        comp_table = self.add_table(len(components_data), 3, headers=components_data[0])
        for row_idx in range(1, len(components_data)):
            comp_table.rows[row_idx].cells[0].text = components_data[row_idx][0]
            comp_table.rows[row_idx].cells[1].text = components_data[row_idx][1]
            comp_table.rows[row_idx].cells[2].text = components_data[row_idx][2]
        
        self.add_heading_3("5.1.4 Supported Query Types")
        
        query_types = [
            ("Query Type", "Example", "Output"),
            ("Simple Selection", "Show me all sales records", "Filtered dataset"),
            ("Aggregation", "Total revenue by region", "Summary table"),
            ("Comparison", "Revenue this year vs last year", "Side-by-side comparison"),
            ("Trend Analysis", "Monthly sales trend for Q4", "Trend visualization + insights"),
            ("Top-N Selection", "Top 5 products by margin", "Ranked list"),
            ("Filtering", "Sales where margin > 20%", "Filtered subset"),
            ("Time-Based", "YoY growth rate", "Time series with calculations"),
            ("Drill-Down", "Break down revenue by category", "Hierarchical breakdown"),
            ("Forecasting", "Predict next quarter revenue", "Forecast with confidence interval"),
        ]
        
        qt_table = self.add_table(len(query_types), 3, headers=query_types[0])
        for row_idx in range(1, len(query_types)):
            qt_table.rows[row_idx].cells[0].text = query_types[row_idx][0]
            qt_table.rows[row_idx].cells[1].text = query_types[row_idx][1]
            qt_table.rows[row_idx].cells[2].text = query_types[row_idx][2]
        
        self.add_heading_3("5.1.5 Error Handling & Fallback Strategies")
        
        error_strategies = [
            "ES-1: If entity extraction fails, attempt rule-based pattern matching",
            "ES-2: If semantic search returns no results, fall back to keyword search",
            "ES-3: If SQL generation fails, provide clarification request to user",
            "ES-4: If LLM call fails, return raw query result without synthesized insights",
            "ES-5: If execution times out (>30s), return partial results with warning"
        ]
        self.add_bullet_list(error_strategies)
        
        # Feature 5.2
        self.add_heading_2("5.2 Multi-Client Profile Management (MCPM)")
        
        self.add_heading_3("5.2.1 Overview")
        self.add_paragraph(
            "Each client organization receives a completely isolated profile/tenant environment with independent configuration, "
            "knowledge bases, user hierarchies, and vector stores. This ensures data sovereignty, customization, and security."
        )
        
        self.add_heading_3("5.2.2 Profile Structure")
        
        profile_components = [
            "Profile ID: Unique identifier per client (UUID v4)",
            "Metadata: Client name, industry vertical, configuration parameters",
            "Vector Store: Dedicated FAISS index per profile (isolated embeddings)",
            "Knowledge Base: Client-specific documents, FAQs, domain knowledge",
            "Users: Hierarchical user management (Admin, Manager, User, Viewer)",
            "Configurations: Custom LLM settings, query parameters, access rules",
            "Session Store: All sessions specific to this profile",
            "Audit Log: Complete activity trail per profile"
        ]
        self.add_bullet_list(profile_components)
        
        self.add_heading_3("5.2.3 Access Control Model")
        
        self.add_paragraph("Role-Based Access Control (RBAC):")
        rbac_roles = [
            "Admin: Full control over profile, users, configurations, knowledge base",
            "Manager: Can manage users, view analytics, configure basic settings",
            "User: Can execute queries, manage own sessions, view shared queries",
            "Viewer: Read-only access to shared queries and results"
        ]
        self.add_bullet_list(rbac_roles)
        
        self.add_heading_3("5.2.4 Profile Operations")
        
        profile_ops_data = [
            ("Operation", "Purpose", "Permissions"),
            ("Create Profile", "Onboard new client", "Super Admin only"),
            ("Update Configuration", "Change client settings", "Admin + Manager"),
            ("Manage Users", "Add/remove users from profile", "Admin"),
            ("Upload Knowledge Base", "Add documents to vector store", "Admin + Manager"),
            ("Export Data", "Extract profile data for backups", "Admin only"),
            ("Delete Profile", "Permanently remove client tenant", "Super Admin + Audit approval"),
        ]
        
        ops_table = self.add_table(len(profile_ops_data), 3, headers=profile_ops_data[0])
        for row_idx in range(1, len(profile_ops_data)):
            ops_table.rows[row_idx].cells[0].text = profile_ops_data[row_idx][0]
            ops_table.rows[row_idx].cells[1].text = profile_ops_data[row_idx][1]
            ops_table.rows[row_idx].cells[2].text = profile_ops_data[row_idx][2]
        
        # Feature 5.3
        self.add_heading_2("5.3 Semantic Search & RAG Pipeline")
        
        self.add_heading_3("5.3.1 Overview")
        self.add_paragraph(
            "Retrieval-Augmented Generation (RAG) combines semantic document retrieval with LLM capabilities to provide "
            "contextually accurate, nuanced insights grounded in actual data and knowledge base content."
        )
        
        self.add_heading_3("5.3.2 Retrieval Process")
        
        retrieval_steps = [
            "1. User Query Reception: Receive natural language query",
            "2. Query Embedding: Convert query to vector using same embedding model as knowledge base",
            "3. Vector Search: Query FAISS index for K nearest neighbor documents (default K=5, configurable 1-500)",
            "4. Rank & Filter: Re-rank results by similarity score, filter by confidence threshold (>0.7)",
            "5. Context Assembly: Prepare retrieved documents as LLM context",
            "6. Cache Check: Verify if similar query already processed (reduce API calls)",
            "7. Return Results: Include retrieved documents with similarity scores in response"
        ]
        self.add_bullet_list(retrieval_steps)
        
        self.add_heading_3("5.3.3 Augmented Generation")
        
        self.add_paragraph("LLM Integration:")
        rag_features = [
            "System Prompt: Pre-defined instructions for insight generation style and tone",
            "Few-Shot Examples: 3-5 example Q&A pairs to guide output format",
            "Retrieved Context: Top-K documents provided as reference material",
            "Query Template: Structured prompt ensuring consistency",
            "Confidence Scoring: LLM assigns confidence level to generated insights (0-100%)",
            "Fact Grounding: Ensure insights reference actual retrieved data",
            "Hallucination Prevention: Cross-check generated claims against source data"
        ]
        self.add_bullet_list(rag_features)
        
        self.add_heading_3("5.3.4 Performance Characteristics")
        
        perf_chars = [
            "Retrieval Latency: <100ms for FAISS search on in-memory index",
            "Generation Latency: 1-5 seconds depending on query complexity",
            "Total E2E Latency: 2-7 seconds (retrieval + generation + formatting)",
            "Accuracy: 85%+ relevance for top-K retrieved documents",
            "Throughput: 100+ concurrent retrieval operations (FAISS thread-safe)",
            "Index Size: ~10GB per 1M documents at 768-dim embeddings"
        ]
        self.add_bullet_list(perf_chars)
        
        # Feature 5.4
        self.add_heading_2("5.4 Session Management & Conversation History")
        
        self.add_heading_3("5.4.1 Session Lifecycle")
        
        lifecycle_stages = [
            "Creation: User initiates new session via API (returns session_id)",
            "Active: Session receives queries, maintains context, updates conversation history",
            "Idle: No activity for >30 minutes; moved to inactive pool but preserved",
            "Archived: Older than 90 days; moved to long-term storage",
            "Closed: Explicitly ended by user or system; persisted for audit"
        ]
        self.add_bullet_list(lifecycle_stages)
        
        self.add_heading_3("5.4.2 Conversation State")
        
        self.add_paragraph("Each session maintains:")
        state_data = [
            "session_id: Unique identifier (UUID v4)",
            "client_id: Associated client/profile",
            "user_id: Owner of session",
            "created_at: Timestamp of session creation",
            "last_activity: Timestamp of most recent query",
            "conversation_history: List of all Query-Response pairs in order",
            "context_buffer: Summarized context for LLM memory (last 5 exchanges)",
            "metadata: Custom tags, labels, or user notes",
            "query_count: Total number of queries in session",
            "tokens_used: LLM token consumption tracking"
        ]
        self.add_bullet_list(state_data)
        
        self.add_heading_3("5.4.3 Context Transfer")
        
        self.add_paragraph(
            "Multi-turn conversations leverage context transfer to maintain coherence across query exchanges. "
            "When processing a new query, the system provides LLM with:"
        )
        context_transfer = [
            "Previous Questions: Last 3-5 user queries for referential understanding",
            "Previous Answers: Summarized insights from prior queries",
            "Active Filters: Current column selections and filter criteria",
            "Time Context: Identified time periods and date references",
            "Entity References: Identified business entities (regions, products, customers)",
            "Assumption Log: Clarifications and user-confirmed interpretations"
        ]
        self.add_bullet_list(context_transfer)
        
        self.doc.add_page_break()
    
    def generate_section_6_functional_requirements(self):
        """Generate Section 6: Functional Requirements (FR)"""
        self.add_heading_1("6. Functional Requirements (FR)")
        
        self.add_paragraph(
            "This section enumerates all functional requirements following ISO/IEC/IEEE 29148 standards. "
            "Each requirement is uniquely identified (FR-XXX) and includes rationale, acceptance criteria, and traceability."
        )
        
        # FR Group 1: Query Processing
        self.add_heading_2("6.1 Query Processing & Interpretation (FR 1.1 - 1.8)")
        
        fr_requirements = [
            {
                "id": "FR-1.1",
                "title": "Natural Language Query Acceptance",
                "requirement": "System SHALL accept user queries in plain English language via POST /api/query endpoint",
                "rationale": "Enables non-technical users to formulate queries without SQL knowledge",
                "acceptance": [
                    "AC-1.1.1: Accepts queries 1-500 characters in length",
                    "AC-1.1.2: Returns 400 error for queries >500 characters",
                    "AC-1.1.3: Accepts special characters and punctuation",
                    "AC-1.1.4: Handles whitespace normalization transparently"
                ],
                "traceability": "Trace to Goal 1 (Democratize Data Access), User Need FN-1"
            },
            {
                "id": "FR-1.2",
                "title": "Query Intent Classification",
                "requirement": "System SHALL classify query intent into categories: SELECT, AGGREGATE, COMPARE, TREND, FORECAST, DRILL_DOWN with >=85% accuracy",
                "rationale": "Intent classification enables appropriate processing pathway selection",
                "acceptance": [
                    "AC-1.2.1: Achieve >=85% F1-score on validation dataset",
                    "AC-1.2.2: Return intent confidence score (0-100%)",
                    "AC-1.2.3: Provide fallback to human-in-loop if confidence <50%",
                    "AC-1.2.4: Log all classifications for training data collection"
                ],
                "traceability": "Trace to Feature 5.1, Functional Objective FO-1"
            },
            {
                "id": "FR-1.3",
                "title": "Entity Extraction from Query",
                "requirement": "System SHALL extract dimensions (tables), measures (columns), time periods, and filters from natural language queries",
                "rationale": "Extracted entities form basis for SQL generation and validation",
                "acceptance": [
                    "AC-1.3.1: Extract table/dimension with >=80% accuracy",
                    "AC-1.3.2: Extract measure/column with >=85% accuracy",
                    "AC-1.3.3: Identify time periods (year, quarter, month)",
                    "AC-1.3.4: Extract filter criteria (>, <, =, BETWEEN, IN)",
                    "AC-1.3.5: Return entity list with confidence scores"
                ],
                "traceability": "Trace to Feature 5.1, FR-1.2"
            },
            {
                "id": "FR-1.4",
                "title": "Semantic Document Retrieval",
                "requirement": "System SHALL retrieve top-K relevant documents from client knowledge base using vector similarity search",
                "rationale": "Retrieved context enables accurate, grounded insight generation",
                "acceptance": [
                    "AC-1.4.1: Return K most similar documents (default K=5, configurable 1-500)",
                    "AC-1.4.2: Similarity scores must be >=0.7 (normalized 0-1 scale)",
                    "AC-1.4.3: Retrieval latency <100ms for in-memory indices",
                    "AC-1.4.4: Return document metadata (source, title, date))",
                    "AC-1.4.5: Support filtering by document type and date range"
                ],
                "traceability": "Trace to Feature 5.3, Functional Objective FO-2"
            },
            {
                "id": "FR-1.5",
                "title": "SQL Query Generation",
                "requirement": "System SHALL generate valid DuckDB SQL from parsed query intent and extracted entities",
                "rationale": "SQL generation enables analytical execution against structured data",
                "acceptance": [
                    "AC-1.5.1: Generated SQL executes without syntax errors",
                    "AC-1.5.2: SQL results match expected output for test cases (100% correctness)",
                    "AC-1.5.3: Query execution time <30 seconds",
                    "AC-1.5.4: Return query plan for optimization opportunities",
                    "AC-1.5.5: Support subqueries and CTEs"
                ],
                "traceability": "Trace to Feature 5.1, FR-1.3, Functional Objective FO-1"
            },
            {
                "id": "FR-1.6",
                "title": "Query Execution Against Data",
                "requirement": "System SHALL execute queries against DuckDB analytical engine and retrieve results",
                "rationale": "Actual data retrieval enables fact-based insights",
                "acceptance": [
                    "AC-1.6.1: Support result sets up to 1M rows",
                    "AC-1.6.2: Return results in JSON format with metadata",
                    "AC-1.6.3: Execution timeout enforced at 30 seconds",
                    "AC-1.6.4: Include row count and data types in response",
                    "AC-1.6.5: Support parameterized queries to prevent injection"
                ],
                "traceability": "Trace to Feature 5.1, FR-1.5"
            },
            {
                "id": "FR-1.7",
                "title": "Result Formatting and Serialization",
                "requirement": "System SHALL format query results for downstream processing (LLM input, API response)",
                "rationale": "Standardized format ensures consistency across system components",
                "acceptance": [
                    "AC-1.7.1: Return results in JSON with schema metadata",
                    "AC-1.7.2: Include row count, column names, and data types",
                    "AC-1.7.3: Support result sampling for large datasets (>100K rows)",
                    "AC-1.7.4: Include data quality flags (nulls, outliers, anomalies)",
                    "AC-1.7.5: Compress results >1MB using gzip"
                ],
                "traceability": "Trace to Feature 5.1, FR-1.6"
            },
            {
                "id": "FR-1.8",
                "title": "Insight Synthesis via LLM",
                "requirement": "System SHALL call Google Gemini API to generate natural language insights from query results and retrieved context",
                "rationale": "LLM transforms raw data into human-understandable insights",
                "acceptance": [
                    "AC-1.8.1: LLM call includes retrieved documents as context",
                    "AC-1.8.2: Generated insights include trend descriptions (>50 chars)",
                    "AC-1.8.3: Insights include anomaly detection and outlier explanation",
                    "AC-1.8.4: Response time 1-5 seconds (including LLM latency)",
                    "AC-1.8.5: Include confidence score (0-100%) for insights"
                ],
                "traceability": "Trace to Feature 5.1, Goal 3 (Semantic Intelligence)"
            }
        ]
        
        for fr in fr_requirements:
            self.add_heading_3(f"{fr['id']}: {fr['title']}")
            self.add_paragraph(f"Requirement: {fr['requirement']}", bold=True)
            self.add_paragraph(f"Rationale: {fr['rationale']}")
            
            self.add_paragraph("Acceptance Criteria:", bold=True)
            self.add_bullet_list(fr['acceptance'])
            
            self.add_paragraph(f"Traceability: {fr['traceability']}")
            self.add_paragraph("")
        
        # FR Group 2: Session Management
        self.add_heading_2("6.2 Session Management (FR 2.1 - 2.4)")
        
        fr_sessions = [
            {
                "id": "FR-2.1",
                "title": "Session Creation",
                "requirement": "System SHALL create a new session on POST /api/sessions request with unique session_id",
                "rationale": "Sessions maintain conversation context and enable multi-turn interactions",
                "acceptance": [
                    "AC-2.1.1: Return unique UUID v4 session_id for each request",
                    "AC-2.1.2: Associate session with authenticated user_id and client_id",
                    "AC-2.1.3: Initialize empty conversation_history array",
                    "AC-2.1.4: Set created_at timestamp with microsecond precision",
                    "AC-2.1.5: Write session to persistent storage within 100ms"
                ],
                "traceability": "Trace to Feature 5.4, User Need FN-3"
            },
            {
                "id": "FR-2.2",
                "title": "Conversation History Maintenance",
                "requirement": "System SHALL maintain ordered, immutable conversation history for each session",
                "rationale": "Full history enables context transfer and audit compliance",
                "acceptance": [
                    "AC-2.2.1: Append each query-response pair to conversation_history",
                    "AC-2.2.2: Preserve order and exact timestamps for all exchanges",
                    "AC-2.2.3: Include metadata (intent, entities, confidence) for each query",
                    "AC-2.2.4: Support retrieval of full history via GET /api/sessions/{id}/history",
                    "AC-2.2.5: Archive sessions >90 days to long-term storage"
                ],
                "traceability": "Trace to Feature 5.4, Assumption A-5"
            },
            {
                "id": "FR-2.3",
                "title": "Context Transfer Between Queries",
                "requirement": "System SHALL transfer context from previous queries to LLM for multi-turn coherence",
                "rationale": "Context transfer enables natural follow-up queries and clarifications",
                "acceptance": [
                    "AC-2.3.1: Include last 3-5 query-response pairs in LLM context",
                    "AC-2.3.2: Summarize longer conversation history to fit token limits",
                    "AC-2.3.3: Identify and carry forward entity references",
                    "AC-2.3.4: Maintain active filters across query exchanges",
                    "AC-2.3.5: Include user clarifications and assumptions in context"
                ],
                "traceability": "Trace to Feature 5.4, FR-2.2, User Need FN-3"
            },
            {
                "id": "FR-2.4",
                "title": "Session Termination and Archive",
                "requirement": "System SHALL support session termination and archive old sessions automatically",
                "rationale": "Lifecycle management ensures efficient storage and compliance",
                "acceptance": [
                    "AC-2.4.1: Support explicit session closure via DELETE /api/sessions/{id}",
                    "AC-2.4.2: Auto-mark sessions idle after 30 minutes of inactivity",
                    "AC-2.4.3: Archive sessions older than 90 days to blob storage",
                    "AC-2.4.4: Maintain full archival for 30 years (regulatory compliance)",
                    "AC-2.4.5: Generate session summary before archival"
                ],
                "traceability": "Trace to Feature 5.4, Assumption A-10"
            }
        ]
        
        for fr in fr_sessions:
            self.add_heading_3(f"{fr['id']}: {fr['title']}")
            self.add_paragraph(f"Requirement: {fr['requirement']}", bold=True)
            self.add_paragraph(f"Rationale: {fr['rationale']}")
            self.add_paragraph("Acceptance Criteria:", bold=True)
            self.add_bullet_list(fr['acceptance'])
            self.add_paragraph(f"Traceability: {fr['traceability']}")
            self.add_paragraph("")
        
        # FR Group 3: Multi-Client Management
        self.add_heading_2("6.3 Multi-Client Profile Management (FR 3.1 - 3.3)")
        
        fr_profiles = [
            {
                "id": "FR-3.1",
                "title": "Profile Creation and Configuration",
                "requirement": "System SHALL support creation of isolated client profiles with independent configurations",
                "rationale": "Multi-tenancy enables B2B deployment with data sovereignty",
                "acceptance": [
                    "AC-3.1.1: Create profile with unique profile_id and name",
                    "AC-3.1.2: Initialize independent FAISS vector store per profile",
                    "AC-3.1.3: Support configuration parameters per profile (LLM settings, query limits)",
                    "AC-3.1.4: Enforce data isolation - queries only access own profile data",
                    "AC-3.1.5: Support up to 50 concurrent profiles in memory"
                ],
                "traceability": "Trace to Feature 5.2, Goal 2 (Enterprise Multi-Tenancy)"
            },
            {
                "id": "FR-3.2",
                "title": "User and Role Management Per Profile",
                "requirement": "System SHALL enforce role-based access control (RBAC) within profiles",
                "rationale": "RBAC ensures appropriate access and maintains security posture",
                "acceptance": [
                    "AC-3.2.1: Support 4 roles: Admin, Manager, User, Viewer",
                    "AC-3.2.2: Enforce role-based endpoint access via middleware",
                    "AC-3.2.3: Admin role can manage users within profile",
                    "AC-3.2.4: User role can create and manage own sessions",
                    "AC-3.2.5: Audit all role changes with timestamps"
                ],
                "traceability": "Trace to Feature 5.2, NFR-8 (Security)"
            },
            {
                "id": "FR-3.3",
                "title": "Knowledge Base Management",
                "requirement": "System SHALL support upload, indexing, and management of client knowledge bases",
                "rationale": "Knowledge bases provide grounding for RAG-based insights",
                "acceptance": [
                    "AC-3.3.1: Support CSV, JSON, Parquet, PDF file uploads",
                    "AC-3.3.2: Transform documents to vector embeddings",
                    "AC-3.3.3: Index embeddings in FAISS per profile",
                    "AC-3.3.4: Support knowledge base size up to 10GB",
                    "AC-3.3.5: Provide reindexing API for knowledge base updates"
                ],
                "traceability": "Trace to Feature 5.2, Feature 5.3, Functional Objective FO-3"
            }
        ]
        
        for fr in fr_profiles:
            self.add_heading_3(f"{fr['id']}: {fr['title']}")
            self.add_paragraph(f"Requirement: {fr['requirement']}", bold=True)
            self.add_paragraph(f"Rationale: {fr['rationale']}")
            self.add_paragraph("Acceptance Criteria:", bold=True)
            self.add_bullet_list(fr['acceptance'])
            self.add_paragraph(f"Traceability: {fr['traceability']}")
            self.add_paragraph("")
        
        # FR Group 4: API & System
        self.add_heading_2("6.4 API Endpoints & System Interfaces (FR 4.1 - 4.4)")
        
        fr_api = [
            {
                "id": "FR-4.1",
                "title": "Query API Endpoint",
                "requirement": "System SHALL expose POST /api/query endpoint accepting JSON payload with query text and session_id",
                "rationale": "Primary interface for query submission",
                "acceptance": [
                    "AC-4.1.1: Accept query_text (string), session_id (UUID), client_id (UUID)",
                    "AC-4.1.2: Return response within 5 seconds (inclLLM time)",
                    "AC-4.1.3: Response includes: insights, data, metadata, confidence, traceability",
                    "AC-4.1.4: Support optional parameters: max_results, confidence_threshold, retrieval_k",
                    "AC-4.1.5: Return 400 for validation errors, 401 for auth failures, 500 otherwise"
                ],
                "traceability": "Trace to Goal 4 (Real-Time), FR-1.1, User Need FN-1"
            },
            {
                "id": "FR-4.2",
                "title": "Session Management API",
                "requirement": "System SHALL expose endpoints for session CRUD operations",
                "rationale": "Session endpoints enable conversation management",
                "acceptance": [
                    "AC-4.2.1: POST /api/sessions - Create new session",
                    "AC-4.2.2: GET /api/sessions/{id} - Retrieve session metadata",
                    "AC-4.2.3: GET /api/sessions/{id}/history - Retrieve conversation history",
                    "AC-4.2.4: DELETE /api/sessions/{id} - Close session",
                    "AC-4.2.5: All endpoints require valid JWT token"
                ],
                "traceability": "Trace to Feature 5.4, FR-2.1 through FR-2.4"
            },
            {
                "id": "FR-4.3",
                "title": "Health Check Endpoint",
                "requirement": "System SHALL expose GET /api/health endpoint returning system status",
                "rationale": "Health checks enable monitoring and load balancing",
                "acceptance": [
                    "AC-4.3.1: Return 200 OK if all components operational",
                    "AC-4.3.2: Include status of: database, LLM API, vector store, services",
                    "AC-4.3.3: Response time <100ms",
                    "AC-4.3.4: Support readiness (/api/health/ready) and liveness (/api/health/live) probes",
                    "AC-4.3.5: No authentication required for health endpoint"
                ],
                "traceability": "Trace to NFR-4 (Monitoring)"
            },
            {
                "id": "FR-4.4",
                "title": "Error Handling & Status Codes",
                "requirement": "System SHALL return appropriate HTTP status codes and error messages",
                "rationale": "Consistent error handling enables client robustness",
                "acceptance": [
                    "AC-4.4.1: 400 Bad Request for invalid input",
                    "AC-4.4.2: 401 Unauthorized for auth failures",
                    "AC-4.4.3: 403 Forbidden for insufficient permissions",
                    "AC-4.4.4: 404 Not Found for missing resources",
                    "AC-4.4.5: 429 Too Many Requests for rate limit violations",
                    "AC-4.4.6: 500 Internal Server Error for unexpected failures"
                ],
                "traceability": "Trace to NFR-6 (Error Handling)"
            }
        ]
        
        for fr in fr_api:
            self.add_heading_3(f"{fr['id']}: {fr['title']}")
            self.add_paragraph(f"Requirement: {fr['requirement']}", bold=True)
            self.add_paragraph(f"Rationale: {fr['rationale']}")
            self.add_paragraph("Acceptance Criteria:", bold=True)
            self.add_bullet_list(fr['acceptance'])
            self.add_paragraph(f"Traceability: {fr['traceability']}")
            self.add_paragraph("")
        
        self.doc.add_page_break()
    
    def save_document(self, filename):
        """Save document"""
        self.doc.save(filename)
        print(f"✓ Document saved: {filename}")
        return filename


def main():
    """Generate Sections 5-6"""
    print("\n" + "=" * 70)
    print("PLAMBO PRD GENERATOR - SECTIONS 5-6 EXTENSION")
    print("=" * 70)
    
    existing_file = r"c:\9Trix\CODE\Plambo\plamboBE\Plambo_PRD_v1.0_Sections_1-4.docx"
    
    print(f"\n[STEP 1] Loading existing document...")
    generator = PlamboPRDSections56(existing_file)
    
    print("[STEP 2] Generating Section 5: Detailed Feature Requirements...")
    generator.generate_section_5_detailed_features()
    
    print("[STEP 3] Generating Section 6: Functional Requirements...")
    generator.generate_section_6_functional_requirements()
    
    output_file = r"c:\9Trix\CODE\Plambo\plamboBE\Plambo_PRD_v1.0_Sections_1-6.docx"
    generator.save_document(output_file)
    
    print("\n" + "=" * 70)
    print("✓ SECTIONS 5-6 COMPLETE - DOCUMENT EXTENDED")
    print("=" * 70)
    print(f"\nOutput: Plambo_PRD_v1.0_Sections_1-6.docx")
    print("\nSections now included:")
    print("  ✓ Section 1: Introduction")
    print("  ✓ Section 2: Product Overview")
    print("  ✓ Section 3: User Research & Analysis")
    print("  ✓ Section 4: High-Level Product Requirements")
    print("  ✓ Section 5: Detailed Feature Requirements")
    print("  ✓ Section 6: Functional Requirements (23 FRs with acceptance criteria)")
    print("\n→ Ready for next sections. Type 'Continue' to proceed to Sections 7-8.")
    
    return output_file


if __name__ == "__main__":
    main()

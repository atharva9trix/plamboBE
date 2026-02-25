#!/usr/bin/env python3
"""
PRD Generator for Plambo Backend (TatvaAI)
Generates a comprehensive Product Requirements Document in Word format
Following ISO/IEC/IEEE 29148 standards
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading_1(doc, text):
    """Add Heading 1"""
    heading = doc.add_heading(text, level=1)
    heading.style = 'Heading 1'
    return heading

def add_heading_2(doc, text):
    """Add Heading 2"""
    heading = doc.add_heading(text, level=2)
    heading.style = 'Heading 2'
    return heading

def add_heading_3(doc, text):
    """Add Heading 3"""
    heading = doc.add_heading(text, level=3)
    heading.style = 'Heading 3'
    return heading

def add_normal_paragraph(doc, text):
    """Add normal paragraph"""
    return doc.add_paragraph(text)

def add_bold_paragraph(doc, text):
    """Add bold paragraph"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.bold = True
    return p

def add_table(doc, rows, cols):
    """Add table to document"""
    return doc.add_table(rows=rows, cols=cols)

def set_table_header_format(table):
    """Format table header row"""
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        cell._element.get_or_add_tcPr().append(
            doc.element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd',
                                    {'w:fill': '4472C4'}
            )
        )

def create_prd():
    """Generate comprehensive PRD document"""
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_paragraph()
    title_run = title.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    title_run.font.size = Pt(28)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Plambo Conversational Business Intelligence Platform")
    subtitle_run.font.size = Pt(16)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    version = doc.add_paragraph()
    version_run = version.add_run("Version: 1.0")
    version_run.font.size = Pt(12)
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(12)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    add_heading_1(doc, "Table of Contents")
    toc_items = [
        "1. Introduction",
        "   1.1 Purpose",
        "   1.2 Scope",
        "   1.3 Intended Audience",
        "   1.4 Definitions & Acronyms",
        "   1.5 References",
        "2. Product Overview",
        "   2.1 Vision Statement",
        "   2.2 Product Goals",
        "   2.3 Target Users",
        "   2.4 Assumptions",
        "   2.5 Constraints",
        "3. User Research",
        "   3.1 User Personas",
        "   3.2 User Environment",
        "   3.3 User Problems & Needs",
        "4. High-Level Product Requirements",
        "   4.1 Feature Summary",
        "   4.2 Feature Prioritization",
        "5. Detailed Feature Requirements",
        "6. Functional Requirements (FR)",
        "7. Non-Functional Requirements (NFR)",
        "8. System Architecture Overview",
        "9. UI/UX Requirements",
        "10. Data Requirements",
        "11. API Requirements",
        "12. Acceptance Criteria Summary",
        "13. Release Plan & Roadmap",
        "14. Risks & Mitigation",
        "15. Appendices"
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 1: INTRODUCTION
    # ============================================================
    add_heading_1(doc, "1. Introduction")
    
    add_heading_2(doc, "1.1 Purpose")
    add_normal_paragraph(doc, 
        "This Product Requirements Document (PRD) defines the comprehensive requirements for the Plambo Conversational Business Intelligence Platform, "
        "an AI-powered backend system designed to process natural language queries and deliver actionable business insights. The document specifies "
        "functional and non-functional requirements, architectural considerations, and acceptance criteria in compliance with ISO/IEC/IEEE 29148 standards.")
    
    add_heading_2(doc, "1.2 Scope")
    add_normal_paragraph(doc,
        "This PRD covers the Plambo Backend (TatvaAI) system, including:")
    scope_items = [
        "Conversational query processing engine",
        "Multi-client profile management and isolation",
        "Vector-based document retrieval and RAG (Retrieval-Augmented Generation)",
        "Session management and conversation history tracking",
        "Data analysis and insight generation using LLM integration",
        "File upload and processing capabilities",
        "RESTful API endpoints for query, data, and session management",
        "System health monitoring and error handling",
        "Web search integration capabilities"
    ]
    for item in scope_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_2(doc, "1.3 Intended Audience")
    add_normal_paragraph(doc,
        "This document is intended for:")
    audience_items = [
        "Product Managers - for feature planning and prioritization",
        "Software Architects - for understanding system design requirements",
        "Backend Developers - for implementation guidance",
        "QA Engineers - for test case development",
        "DevOps/Infrastructure Teams - for deployment and infrastructure requirements",
        "Business Stakeholders - for understanding product capabilities and roadmap"
    ]
    for item in audience_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_2(doc, "1.4 Definitions & Acronyms")
    acronyms_data = [
        ("RAG", "Retrieval-Augmented Generation"),
        ("LLM", "Large Language Model"),
        ("FAISS", "Facebook AI Similarity Search"),
        ("API", "Application Programming Interface"),
        ("REST", "Representational State Transfer"),
        ("BI", "Business Intelligence"),
        ("DuckDB", "In-process SQL database system"),
        ("TatvaAI", "Core AI engine (Tatva = Essence in Sanskrit)"),
        ("Session", "A user interaction context maintaining conversation history"),
        ("Profile/Client", "A distinct organizational entity with isolated vector stores and knowledge bases"),
        ("Vector Store", "FAISS index for semantic document retrieval"),
        ("NFR", "Non-Functional Requirement"),
        ("FR", "Functional Requirement"),
        ("SLA", "Service Level Agreement"),
        ("QA", "Quality Assurance")
    ]
    
    table = add_table(doc, len(acronyms_data) + 1, 2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = "Term"
    header_cells[1].text = "Definition"
    for i, (term, defn) in enumerate(acronyms_data, 1):
        table.rows[i].cells[0].text = term
        table.rows[i].cells[1].text = defn
    
    add_heading_2(doc, "1.5 References")
    add_normal_paragraph(doc, "Relevant standards and guidelines:")
    references = [
        "ISO/IEC/IEEE 29148:2018 - Systems and software engineering requirements specification",
        "Flask Documentation - Web framework specifications",
        "FAISS Documentation - Vector store implementation",
        "OpenAI/Google Gemini API Documentation - LLM integration reference",
        "RESTful API Design Best Practices"
    ]
    for ref in references:
        doc.add_paragraph(ref, style='List Bullet')
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 2: PRODUCT OVERVIEW
    # ============================================================
    add_heading_1(doc, "2. Product Overview")
    
    add_heading_2(doc, "2.1 Vision Statement")
    add_normal_paragraph(doc,
        "The Plambo Conversational Business Intelligence Platform aims to democratize data-driven decision-making by enabling business users "
        "to interact with their organizational data through natural language conversations. By leveraging advanced AI, semantic search, and "
        "LLM technologies, the platform transforms complex data queries into accessible, intuitive interactions while maintaining privacy, "
        "security, and multi-tenancy isolation for enterprise clients.")
    
    add_heading_2(doc, "2.2 Product Goals")
    goals = [
        ("Enable Natural Language Queries", 
         "Allow users to ask business questions in plain English without requiring SQL or technical expertise"),
        ("Deliver Contextual Insights",
         "Provide analyzed, actionable insights rather than raw data through AI-powered interpretation"),
        ("Support Multi-Tenant Architecture",
         "Maintain strict data isolation across different clients while operating from a single backend"),
        ("Maintain 99.5% Availability",
         "Ensure reliable service with minimal downtime for business-critical operations"),
        ("Scale to Multiple Data Sources",
         "Support diverse data formats (Parquet, CSV, SQL databases) unified through vector semantics"),
        ("Preserve Conversation Context",
         "Enable multi-turn conversations with session-based memory for coherent interactions"),
        ("Ensure Enterprise Security",
         "Implement authorization, data privacy, and compliance controls at all layers")
    ]
    
    for goal_title, goal_desc in goals:
        p = doc.add_paragraph()
        p.add_run(goal_title + ": ").bold = True
        p.add_run(goal_desc)
    
    add_heading_2(doc, "2.3 Target Users")
    users = [
        ("Business Analysts", "Query data to identify trends, anomalies, and business metrics without SQL expertise"),
        ("Executives & Decision Makers", "Access high-level insights and KPIs through conversational interface"),
        ("Data Scientists", "Leverage API for programmatic access to analysis and data exploration"),
        ("Product Managers", "Monitor product metrics and user engagement through natural language queries"),
        ("Operations Teams", "Track performance indicators and system health through query interface"),
        ("Enterprise Administrators", "Manage multiple client profiles and ensure security/compliance")
    ]
    
    for user_type, user_desc in users:
        p = doc.add_paragraph()
        p.add_run(user_type + ": ").bold = True
        p.add_run(user_desc)
    
    add_heading_2(doc, "2.4 Assumptions")
    assumptions = [
        "Users have basic familiarity with business terminology and their organizational structure",
        "Data is pre-processed and stored in supported formats (Parquet, accessible databases)",
        "Network connectivity to the backend service is available and stable",
        "Users are authenticated and authorized to access their respective client profiles",
        "LLM services (Ollama local or Google Gemini API) are available and functional",
        "Vector stores are pre-built and indexed for supported profiles",
        "Database backend (PostgreSQL/DuckDB) is properly configured and accessible",
        "Session data should be retained for at least 30 days for conversation history"
    ]
    for assumption in assumptions:
        doc.add_paragraph(assumption, style='List Bullet')
    
    add_heading_2(doc, "2.5 Constraints")
    constraints = [
        ("Technical", "System operates on Flask framework limited to Python ecosystem; vector store operations are CPU-intensive"),
        ("Operational", "LLM token limits (API calls) may restrict query volume without rate limiting implementation"),
        ("Compliance", "Multi-tenant isolation requires strict authentication to prevent cross-client data access"),
        ("Performance", "Large vector store searches may introduce latency for complex queries"),
        ("Data", "Requires pre-processed, structured data; unstructured data ingestion not supported in v1.0"),
        ("Scale", "Initial deployment targets 5-10 enterprise clients; scaling may require architecture changes")
    ]
    
    for constraint_type, constraint_desc in constraints:
        p = doc.add_paragraph()
        p.add_run(f"{constraint_type}: ").bold = True
        p.add_run(constraint_desc)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 3: USER RESEARCH
    # ============================================================
    add_heading_1(doc, "3. User Research")
    
    add_heading_2(doc, "3.1 User Personas")
    
    personas = [
        {
            "name": "Alex Chen - Business Analyst",
            "description": "5 years in business analytics, comfortable with Excel, limited SQL knowledge",
            "goals": "Quickly access sales trends and perform ad-hoc analysis without requesting IT support",
            "pain_points": "Complex SQL queries, time-consuming data warehouse access, dependency on data team",
            "technical_skill": "Medium"
        },
        {
            "name": "Priya Sharma - Executive Manager",
            "description": "Senior leader responsible for P&L, needs executive dashboards and quick insights",
            "goals": "Access real-time business metrics and competitive intelligence to inform decisions",
            "pain_points": "Delayed reporting cycles, inability to ask follow-up questions to analytics team",
            "technical_skill": "Low"
        },
        {
            "name": "Marco Rodriguez - Data Engineer",
            "description": "Technical practitioner building data pipelines, familiar with APIs and automation",
            "goals": "Integrate BI platform into automated workflows and build custom applications on top",
            "pain_points": "Limited API capabilities, rigid output formats, lack of programmatic access",
            "technical_skill": "High"
        },
        {
            "name": "Lisa Thompson - IT Administrator",
            "description": "Responsible for security, compliance, and multi-tenant infrastructure",
            "goals": "Ensure secure isolation between clients, maintain audit trails, comply with regulations",
            "pain_points": "Data leakage risks, complexity of multi-tenant management, audit visibility",
            "technical_skill": "High"
        }
    ]
    
    for persona in personas:
        add_heading_3(doc, persona["name"])
        p = doc.add_paragraph()
        p.add_run("Description: ").bold = True
        p.add_run(persona["description"])
        
        p = doc.add_paragraph()
        p.add_run("Goals: ").bold = True
        p.add_run(persona["goals"])
        
        p = doc.add_paragraph()
        p.add_run("Pain Points: ").bold = True
        p.add_run(persona["pain_points"])
        
        p = doc.add_paragraph()
        p.add_run("Technical Skill Level: ").bold = True
        p.add_run(persona["technical_skill"])
    
    add_heading_2(doc, "3.2 User Environment")
    add_normal_paragraph(doc, "Users access the Plambo platform through:")
    environments = [
        "Web-based frontend (to be developed) - primary interface for business users",
        "RESTful API - for programmatic access and integration with third-party tools",
        "Direct SQL/Data exploration - for advanced users and data engineers",
        "Mobile access - future consideration for executive dashboards",
        "Integrated business tools - Slack, Teams, Microsoft Excel (future)"
    ]
    for env in environments:
        doc.add_paragraph(env, style='List Bullet')
    
    add_normal_paragraph(doc, "Infrastructure environment:")
    infra_items = [
        "Cloud-based or on-premises deployment",
        "Integration with existing enterprise databases",
        "Connection to LLM services (Ollama locally or cloud-based)",
        "Vector database storage (FAISS local or cloud-based)",
        "Authentication via SSO or enterprise IAM systems"
    ]
    for item in infra_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_2(doc, "3.3 User Problems & Needs")
    
    problems_table = add_table(doc, 1, 3)
    problems_table.style = 'Light Grid Accent 1'
    header = problems_table.rows[0].cells
    header[0].text = "User Problem"
    header[1].text = "Business Impact"
    header[2].text = "Solution Provided"
    
    problems_data = [
        ("Time-consuming manual data queries", "Days spent on analytics, delayed decisions", "Natural language interface enables instant queries"),
        ("Dependency on data teams", "Bottlenecks, limited accessibility", "Self-service analytics reduces dependency"),
        ("Complex SQL learning curve", "High barrier to entry", "Conversational interface requires no SQL"),
        ("Cross-client data leakage risk", "Compliance violations, security breach", "Strict multi-tenant isolation"),
        ("Limited conversation context", "Repetitive context provision", "Session memory preserves conversation history"),
        ("Limited data source integration", "Siloed data analysis", "Multi-format data ingestion capability")
    ]
    
    for problem, impact, solution in problems_data:
        row = problems_table.add_row().cells
        row[0].text = problem
        row[1].text = impact
        row[2].text = solution
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 4: HIGH-LEVEL REQUIREMENTS
    # ============================================================
    add_heading_1(doc, "4. High-Level Product Requirements")
    
    add_heading_2(doc, "4.1 Feature Summary")
    add_normal_paragraph(doc, "The Plambo platform provides the following core features:")
    
    features = [
        ("Conversational Query Processing", "Process natural language questions and convert to structured analysis"),
        ("Multi-Client Profile Management", "Maintain isolated knowledge bases and vector stores per client"),
        ("Session Management", "Track conversation history and context across multiple user interactions"),
        ("Semantic Document Retrieval", "Search and retrieve relevant documents using vector similarity"),
        ("AI-Powered Insight Generation", "Leverage LLM to analyze data and generate human-readable insights"),
        ("File Upload & Processing", "Accept Parquet and data files for analysis and indexing"),
        ("Health Monitoring", "System health checks and performance monitoring endpoints"),
        ("Web Search Integration", "Augment internal knowledge with web search results when needed"),
        ("API-First Architecture", "Comprehensive REST API for programmatic access"),
        ("Data Description & Validation", "Automatic data type detection and validation")
    ]
    
    for feature_name, feature_desc in features:
        p = doc.add_paragraph()
        p.add_run(feature_name + ": ").bold = True
        p.add_run(feature_desc)
    
    add_heading_2(doc, "4.2 Feature Prioritization (MoSCoW)")
    
    priority_table = add_table(doc, 1, 3)
    priority_table.style = 'Light Grid Accent 1'
    header = priority_table.rows[0].cells
    header[0].text = "Priority"
    header[1].text = "Features"
    header[2].text = "Target Release"
    
    priorities_data = [
        ("MUST", "Query processing, Profile management, Session tracking, API endpoints", "v1.0 (Current)"),
        ("SHOULD", "Web search integration, Health monitoring, File upload/processing", "v1.1"),
        ("COULD", "Advanced analytics, Custom report generation, Mobile API", "v1.2+"),
        ("WON'T", "Real-time data streaming, Graph database integration, Native mobile apps", "Post-v2.0")
    ]
    
    for priority, features_list, release in priorities_data:
        row = priority_table.add_row().cells
        row[0].text = priority
        row[1].text = features_list
        row[2].text = release
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 5: DETAILED FEATURE REQUIREMENTS
    # ============================================================
    add_heading_1(doc, "5. Detailed Feature Requirements")
    
    # Feature 1: Query Processing
    add_heading_2(doc, "5.1 Conversational Query Processing")
    
    add_bold_paragraph(doc, "Feature ID: FEAT-001")
    add_bold_paragraph(doc, "Business Justification")
    add_normal_paragraph(doc,
        "Natural language query processing is the core differentiator enabling non-technical users to interact with data "
        "without SQL knowledge. This feature directly addresses user pain point of data accessibility and reduces dependency on data teams.")
    
    add_bold_paragraph(doc, "User Story")
    add_normal_paragraph(doc,
        "As a Business Analyst, I want to ask questions in English like 'What were our top 5 products by revenue in Q3?' "
        "so that I can quickly get insights without writing SQL or waiting for the data team.")
    
    add_bold_paragraph(doc, "Description")
    add_normal_paragraph(doc,
        "The system accepts natural language queries via REST API endpoint (/api/query) and processes them through "
        "the following pipeline: (1) Query validation and sanitization, (2) Vector similarity search against client knowledge base, "
        "(3) LLM-based query interpretation, (4) Insight generation with supporting context.")
    
    add_bold_paragraph(doc, "Acceptance Criteria")
    criteria = [
        "System accepts POST requests with valid JSON payload containing client_id, query, and conversation_context",
        "Query response time is < 5 seconds for 95% of requests (< 10s for 99%)",
        "System retrieves top-5 relevant documents from vector store based on semantic similarity",
        "LLM-generated responses are accurate and contextually relevant",
        "System returns structured JSON response with query, answer, retrieved_context, and metadata",
        "Invalid queries are rejected with appropriate error messages",
        "Query processing is isolated per client; cross-client data access is impossible"
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    add_bold_paragraph(doc, "Functional Requirements")
    fr_items = [
        "Query validation must reject empty queries and queries exceeding 1000 characters",
        "Vector store retrieval must use top_k=5 by default (configurable)",
        "LLM processing must include generated insights and supporting evidence from retrieved documents",
        "Response must include metadata: query_id, timestamp, client_id, response_time_ms"
    ]
    for item in fr_items:
        doc.add_paragraph("FR-001-" + str(fr_items.index(item) + 1) + ": " + item, style='List Bullet')
    
    add_bold_paragraph(doc, "Non-Functional Requirements")
    nfr_items = [
        "Response latency < 5 seconds for 95% of queries",
        "Query endpoint must handle 100 req/sec sustained load",
        "System must support concurrent users: minimum 50 simultaneous active sessions"
    ]
    for item in nfr_items:
        doc.add_paragraph("NFR-001-" + str(nfr_items.index(item) + 1) + ": " + item, style='List Bullet')
    
    add_bold_paragraph(doc, "Dependencies")
    deps = [
        "FAISS Vector Store must be pre-indexed and available",
        "LLM service (Ollama/Gemini) must be accessible",
        "Profile manager must be initialized with client data"
    ]
    for dep in deps:
        doc.add_paragraph(dep, style='List Bullet')
    
    add_bold_paragraph(doc, "Edge Cases")
    edges = [
        "Empty query string → Return validation error",
        "No matching documents found → System returns fallback response with LLM disclaimer",
        "LLM service timeout → Return partial response with available context",
        "Malformed JSON payload → Return 400 Bad Request",
        "Unauthorized client_id → Return 403 Forbidden"
    ]
    for edge in edges:
        doc.add_paragraph(edge, style='List Bullet')
    
    doc.add_page_break()
    
    # Feature 2: Session Management
    add_heading_2(doc, "5.2 Session Management & Conversation History")
    
    add_bold_paragraph(doc, "Feature ID: FEAT-002")
    add_bold_paragraph(doc, "Business Justification")
    add_normal_paragraph(doc,
        "Session management enables multi-turn conversations where context is preserved across interactions. "
        "This improves user experience and allows for follow-up questions without repetition.")
    
    add_bold_paragraph(doc, "User Story")
    add_normal_paragraph(doc,
        "As an executive, I want to have a conversation where I ask 'What were our Q3 sales?' and then follow up with "
        "'Which region contributed the most?' without repeating context, so I can have natural back-and-forth dialogue.")
    
    add_bold_paragraph(doc, "Description")
    add_normal_paragraph(doc,
        "System manages user sessions with unique session IDs. Each session maintains a conversation history "
        "containing previous questions and answers. New queries are processed with awareness of conversation context.")
    
    add_bold_paragraph(doc, "Acceptance Criteria")
    criteria = [
        "System creates unique session ID on /api/create_session (GET) with user_id parameter",
        "Session data persists in database for minimum 30 days",
        "Each session stores up to 100 previous Q&A pairs in FIFO order",
        "System retrieves last 5 Q&A pairs as conversation context for LLM processing",
        "Session history is accessible only to the user who created the session",
        "Sessions automatically expire after 90 days of inactivity"
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    add_bold_paragraph(doc, "Dependencies")
    deps = [
        "Database backend (PostgreSQL/DuckDB) for session storage",
        "User authentication system for session isolation",
        "Query processing engine (FEAT-001)"
    ]
    for dep in deps:
        doc.add_paragraph(dep, style='List Bullet')
    
    doc.add_page_break()
    
    # Feature 3: Multi-Client Profile Management
    add_heading_2(doc, "5.3 Multi-Client Profile Management")
    
    add_bold_paragraph(doc, "Feature ID: FEAT-003")
    add_bold_paragraph(doc, "Business Justification")
    add_normal_paragraph(doc,
        "Enterprise clients require strict data isolation. Profile management ensures multiple organizations "
        "can share a single backend while maintaining complete data separation and security.")
    
    add_bold_paragraph(doc, "User Story")
    add_normal_paragraph(doc,
        "As an IT administrator, I want to configure different knowledge bases and vector stores for different clients "
        "so that my company's data never mixes with other organizations' data on the shared platform.")
    
    add_bold_paragraph(doc, "Description")
    add_normal_paragraph(doc,
        "System supports multiple client profiles, each with isolated vector stores, knowledge bases, "
        "and configuration. Profile selection is mandatory for query processing.")
    
    add_bold_paragraph(doc, "Acceptance Criteria")
    criteria = [
        "System maintains separate FAISS vector indices per client",
        "Query processing validates client authorization before data access",
        "Client data is never accessible across profile boundaries",
        "/api/clients endpoint returns list of available clients for authenticated user",
        "Invalid client_id returns 403 Forbidden error",
        "System supports minimum 10 concurrent client profiles"
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    doc.add_page_break()
    
    # Feature 4: File Upload & Processing
    add_heading_2(doc, "5.4 File Upload & Data Processing")
    
    add_bold_paragraph(doc, "Feature ID: FEAT-004")
    add_bold_paragraph(doc, "Business Justification")
    add_normal_paragraph(doc,
        "Users need ability to upload and process new data files. This enables ad-hoc analysis on user-provided datasets "
        "and allows integration with various data sources.")
    
    add_bold_paragraph(doc, "User Story")
    add_normal_paragraph(doc,
        "As a data analyst, I want to upload a CSV or Parquet file and have the system automatically detect data types "
        "and make it available for querying so I can analyze new datasets quickly.")
    
    add_bold_paragraph(doc, "Description")
    add_normal_paragraph(doc,
        "System provides endpoint /api/fetch_attribute/dtype/v1.0 for file uploads. Files are processed to extract "
        "data types, columns, and metadata.")
    
    add_bold_paragraph(doc, "Acceptance Criteria")
    criteria = [
        "System accepts Parquet and CSV file uploads via multipart/form-data",
        "Maximum file size supported: 500MB",
        "System automatically detects data types for all columns",
        "Processing completes within 30 seconds for files < 100MB",
        "System returns structured response with data types and column information",
        "Invalid files return appropriate error messages with guidance"
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 6: FUNCTIONAL REQUIREMENTS
    # ============================================================
    add_heading_1(doc, "6. Functional Requirements (FR-xx Format)")
    
    fr_detailed = [
        {
            "id": "FR-001",
            "name": "Query Validation & Sanitization",
            "description": "All incoming queries must be validated for format, length, and content security threats",
            "conditions": "For all POST /api/query requests"
        },
        {
            "id": "FR-002",
            "name": "Vector Store Retrieval",
            "description": "System must retrieve top-K semantically similar documents from FAISS index",
            "conditions": "For each query processing request; K=5 by default"
        },
        {
            "id": "FR-003",
            "name": "LLM Query Processing",
            "description": "Retrieved documents are passed to LLM with user query for insight generation",
            "conditions": "When LLM service is available; fallback if unavailable"
        },
        {
            "id": "FR-004",
            "name": "Multi-Tenant Isolation",
            "description": "Each request must verify client authorization; prevent cross-client data access",
            "conditions": "Mandatory for all data operations"
        },
        {
            "id": "FR-005",
            "name": "Session Creation & Management",
            "description": "System creates unique session IDs; maintains conversation history per session",
            "conditions": "On /api/create_session endpoint call with valid user_id"
        },
        {
            "id": "FR-006",
            "name": "Conversation Context Preservation",
            "description": "Last N (configurable, default=5) Q&A pairs included in LLM context window",
            "conditions": "When processing new query in existing session"
        },
        {
            "id": "FR-007",
            "name": "Health Status Monitoring",
            "description": "System provides health check endpoint returning service status details",
            "conditions": "On /api/health GET request"
        },
        {
            "id": "FR-008",
            "name": "Error Handling & Reporting",
            "description": "All errors return standardized JSON response with error codes and messages",
            "conditions": "On any error condition"
        },
        {
            "id": "FR-009",
            "name": "Web Search Integration",
            "description": "System can optionally augment responses with web search results",
            "conditions": "When web search flag is enabled in query request"
        },
        {
            "id": "FR-010",
            "name": "Response Formatting",
            "description": "All responses return standardized JSON with query_id, timestamp, answer, metadata",
            "conditions": "For all successful responses"
        }
    ]
    
    for fr in fr_detailed:
        add_heading_2(doc, f"{fr['id']}: {fr['name']}")
        p = doc.add_paragraph()
        p.add_run("Description: ").bold = True
        p.add_run(fr['description'])
        
        p = doc.add_paragraph()
        p.add_run("Conditions: ").bold = True
        p.add_run(fr['conditions'])
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 7: NON-FUNCTIONAL REQUIREMENTS
    # ============================================================
    add_heading_1(doc, "7. Non-Functional Requirements")
    
    add_heading_2(doc, "7.1 Performance")
    nfr_perf = [
        ("Response Latency", "95th percentile < 5 seconds, 99th percentile < 10 seconds for query processing"),
        ("Throughput", "System must handle 100 requests/second sustained load"),
        ("Concurrent Sessions", "Support minimum 50 simultaneous active user sessions"),
        ("Vector Search Speed", "FAISS index search < 500ms for top-5 retrieval"),
        ("File Processing", "Process 100MB file within 30 seconds"),
        ("Database Queries", "Session lookup < 100ms")
    ]
    for metric, requirement in nfr_perf:
        p = doc.add_paragraph()
        p.add_run(metric + ": ").bold = True
        p.add_run(requirement)
    
    add_heading_2(doc, "7.2 Availability & Reliability")
    nfr_avail = [
        ("System Uptime", "99.5% availability SLA (maximum 3.6 hours downtime per month)"),
        ("Data Durability", "Session data backed up minimum once daily"),
        ("Recovery Time Objective (RTO)", "Service recovery within 30 minutes of failure"),
        ("Recovery Point Objective (RPO)", "Maximum 1 hour of data loss in disaster scenario"),
        ("Failover", "Automatic failover to secondary instance within 5 minutes")
    ]
    for metric, requirement in nfr_avail:
        p = doc.add_paragraph()
        p.add_run(metric + ": ").bold = True
        p.add_run(requirement)
    
    add_heading_2(doc, "7.3 Security")
    nfr_sec = [
        ("Authentication", "All API endpoints require bearer token authentication (JWT)"),
        ("Authorization", "Client isolation at database and application layers"),
        ("Data Encryption", "All data in transit encrypted via TLS 1.2+"),
        ("Data at Rest", "Sensitive data encrypted using AES-256"),
        ("SQL Injection", "All database queries use parameterized statements"),
        ("Rate Limiting", "API endpoints enforce 1000 requests/hour per authenticated user"),
        ("Audit Logging", "All data access logged with user_id, timestamp, query_id, client_id")
    ]
    for metric, requirement in nfr_sec:
        p = doc.add_paragraph()
        p.add_run(metric + ": ").bold = True
        p.add_run(requirement)
    
    add_heading_2(doc, "7.4 Scalability")
    nfr_scale = [
        ("Horizontal Scaling", "Stateless API design allows horizontal scaling via load balancers"),
        ("Vertical Scaling", "Support for instance sizes from 2GB to 32GB RAM"),
        ("Data Scaling", "Vector stores scalable to 100K+ documents per client"),
        ("User Scaling", "Architecture supports 1000+ concurrent users with proper resource allocation")
    ]
    for metric, requirement in nfr_scale:
        p = doc.add_paragraph()
        p.add_run(metric + ": ").bold = True
        p.add_run(requirement)
    
    add_heading_2(doc, "7.5 Maintainability")
    nfr_maint = [
        ("Code Quality", "Minimum 80% code coverage with automated tests"),
        ("Documentation", "All API endpoints documented with OpenAPI/Swagger specs"),
        ("Logging", "Structured logging with DEBUG, INFO, WARN, ERROR levels"),
        ("Monitoring", "Real-time monitoring dashboards for key metrics"),
        ("Deployment", "Containerized deployment via Docker for consistency")
    ]
    for metric, requirement in nfr_maint:
        p = doc.add_paragraph()
        p.add_run(metric + ": ").bold = True
        p.add_run(requirement)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 8: SYSTEM ARCHITECTURE
    # ============================================================
    add_heading_1(doc, "8. System Architecture Overview")
    
    add_heading_2(doc, "8.1 High-Level Architecture Description")
    add_normal_paragraph(doc,
        "The Plambo platform follows a modular, multi-layered architecture designed for scalability, security, and maintainability:")
    
    add_heading_3(doc, "Architecture Layers")
    
    layers = [
        {
            "name": "API Layer",
            "components": ["Flask REST endpoints", "Request validation", "Response formatting"],
            "description": "Exposes RESTful API for query processing, session management, file handling, and system health monitoring"
        },
        {
            "name": "Business Logic Layer",
            "components": ["QueryService", "SessionManager", "ProfileManager", "LLM Processor"],
            "description": "Core logic for processing queries, managing sessions, handling multi-client profiles, and LLM integration"
        },
        {
            "name": "Data Access Layer",
            "components": ["Database Engine", "Vector Store Interface", "File Handler"],
            "description": "Manages interactions with various data stores including PostgreSQL, DuckDB, FAISS vector stores"
        },
        {
            "name": "Integration Layer",
            "components": ["LLM Connectors", "Web Search API", "External Services"],
            "description": "Integrations with LLM services, web search APIs, and other external services"
        },
        {
            "name": "Infrastructure Layer",
            "components": ["Docker containers", "Load balancers", "Database backups", "Logging/Monitoring"],
            "description": "Deployment and operational infrastructure"
        }
    ]
    
    for layer in layers:
        add_heading_3(doc, layer["name"])
        p = doc.add_paragraph()
        p.add_run("Components: ").bold = True
        p.add_run(", ".join(layer["components"]))
        
        p = doc.add_paragraph()
        p.add_run("Purpose: ").bold = True
        p.add_run(layer["description"])
    
    add_heading_3(doc, "Data Flow")
    add_normal_paragraph(doc,
        "Typical query processing flow: (1) User submits natural language query via API → "
        "(2) Validation & sanitization → (3) Session context retrieval → (4) Vector store semantic search → "
        "(5) LLM inference with context → (6) Response formatting → (7) Logging & response delivery")
    
    add_heading_2(doc, "8.2 Technology Stack")
    
    tech_table = add_table(doc, 1, 3)
    tech_table.style = 'Light Grid Accent 1'
    header = tech_table.rows[0].cells
    header[0].text = "Component"
    header[1].text = "Technology"
    header[2].text = "Justification"
    
    tech_data = [
        ("Web Framework", "Flask (Python)", "Lightweight, extensive ecosystem, rapid development"),
        ("Language", "Python 3.13+", "Rich ML/data science libraries, readable code, strong AI community"),
        ("Vector Search", "FAISS (Facebook AI Similarity Search)", "High-performance, open-source, proven at scale"),
        ("LLM Integration", "Ollama (local) + Google Gemini API (cloud)", "Flexibility, cost-effective local models, advanced cloud models"),
        ("Relational Database", "PostgreSQL / DuckDB", "Robust session storage, in-memory analytics capacity"),
        ("Data Format", "Apache Parquet", "Columnar storage, compression, compatibility with Python/ML tools"),
        ("Containerization", "Docker", "Consistency, reproducibility, cloud deployment"),
        ("API Documentation", "OpenAPI/Swagger", "Interactive API docs, SDK generation"),
        ("Monitoring", "Prometheus + Grafana", "Open-source, scalable monitoring"),
        ("Logging", "Python logging + ELK", "Structured logs, searchable history")
    ]
    
    for component, technology, justification in tech_data:
        row = tech_table.add_row().cells
        row[0].text = component
        row[1].text = technology
        row[2].text = justification
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 9: UI/UX REQUIREMENTS
    # ============================================================
    add_heading_1(doc, "9. UI/UX Requirements")
    
    add_heading_2(doc, "9.1 User Flows")
    
    add_heading_3(doc, "Flow 1: Natural Language Query")
    flow1_steps = [
        "User logs into web interface",
        "Selects client/profile from dropdown",
        "Enters natural language query in text input",
        "System processes query in background",
        "Results displayed with supporting context and source documents",
        "User can view, export, or ask follow-up question"
    ]
    for i, step in enumerate(flow1_steps, 1):
        doc.add_paragraph(f"Step {i}: {step}", style='List Number')
    
    add_heading_3(doc, "Flow 2: Session Management")
    flow2_steps = [
        "User creates new session on login",
        "Conversation history displayed in sidebar (last 10 Q&A pairs)",
        "Each Q&A pair is clickable to view full conversation context",
        "User can create new session or resume existing session",
        "Session metadata (created date, last accessed, query count) displayed"
    ]
    for i, step in enumerate(flow2_steps, 1):
        doc.add_paragraph(f"Step {i}: {step}", style='List Number')
    
    add_heading_3(doc, "Flow 3: File Upload & Analysis")
    flow3_steps = [
        "User navigates to 'Upload Data' section",
        "Selects file from local storage (Parquet or CSV)",
        "System displays data preview and detected data types",
        "User confirms upload; system processes file",
        "File becomes available for querying",
        "Results display analysis based on uploaded data"
    ]
    for i, step in enumerate(flow3_steps, 1):
        doc.add_paragraph(f"Step {i}: {step}", style='List Number')
    
    add_heading_2(doc, "9.2 Wireframes (Textual Description)")
    
    add_heading_3(doc, "Main Dashboard Layout")
    add_normal_paragraph(doc,
        "Header: Logo | User Profile | Logout\n"
        "Left Sidebar: Client selector dropdown | Session history | File uploads | Settings\n"
        "Main Content: Query input text area | Submit button | Previous results | Export button\n"
        "Right Sidebar: Conversation context | Retrieved documents | Settings panel\n"
        "Footer: API status | Help | Documentation links"
    )
    
    add_heading_3(doc, "Query Results View")
    add_normal_paragraph(doc,
        "Top: Original query display | Query execution time\n"
        "Middle: AI-generated answer in clear, readable format\n"
        "Below: Source documents display (list of retrieved documents with confidence scores)\n"
        "Bottom: Follow-up suggestion options | Export/Share buttons\n"
        "Collapsible panels: Raw response data | Query metadata"
    )
    
    add_heading_2(doc, "9.3 Design Guidelines")
    
    guidelines = [
        ("Color Scheme", "Professional blue/white primary, green for success/positive, red for errors"),
        ("Typography", "Sans-serif font (Roboto/Inter), consistent sizing hierarchy"),
        ("Accessibility", "WCAG 2.1 AA compliance, keyboard navigation support, screen reader compatible"),
        ("Responsiveness", "Mobile-first design, responsive breakpoints at 768px, 1024px, 1440px"),
        ("Loading States", "Clear loading indicators, skeleton screens for data, estimated wait time"),
        ("Error Handling", "User-friendly error messages with troubleshooting suggestions"),
        ("Consistency", "Standardized component library, consistent spacing/margins throughout")
    ]
    
    for guideline, details in guidelines:
        p = doc.add_paragraph()
        p.add_run(guideline + ": ").bold = True
        p.add_run(details)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 10: DATA REQUIREMENTS
    # ============================================================
    add_heading_1(doc, "10. Data Requirements")
    
    add_heading_2(doc, "10.1 Data Entities & Structure")
    
    data_entities = [
        {
            "name": "User",
            "attributes": ["user_id (PK)", "email", "name", "organization", "created_date", "last_login"],
            "description": "Represents system users and their metadata"
        },
        {
            "name": "Session",
            "attributes": ["session_id (PK)", "user_id (FK)", "client_id", "created_date", "last_accessed", "status", "conversation_history"],
            "description": "Tracks user sessions and conversation history"
        },
        {
            "name": "Query",
            "attributes": ["query_id (PK)", "session_id (FK)", "query_text", "timestamp", "response", "execution_time_ms", "document_count"],
            "description": "Records all queries and responses for audit and analytics"
        },
        {
            "name": "Document",
            "attributes": ["doc_id (PK)", "client_id", "content", "embedding_vector", "metadata", "source", "created_date"],
            "description": "Documents indexed in vector stores for semantic search"
        },
        {
            "name": "Client/Profile",
            "attributes": ["client_id (PK)", "name", "data_dir", "vector_store_path", "config_json", "created_date"],
            "description": "Multi-tenant configuration for different client profiles"
        }
    ]
    
    for entity in data_entities:
        add_heading_3(doc, entity["name"])
        p = doc.add_paragraph()
        p.add_run("Attributes: ").bold = True
        p.add_run(", ".join(entity["attributes"]))
        
        p = doc.add_paragraph()
        p.add_run("Description: ").bold = True
        p.add_run(entity["description"])
    
    add_heading_2(doc, "10.2 Data Validation Rules")
    
    validation_rules = [
        ("Query Length", "Minimum 3 characters, maximum 1000 characters"),
        ("Query Format", "Must be non-empty string; SQL injection attempts blocked"),
        ("Client ID Format", "Alphanumeric, underscores allowed, case-sensitive"),
        ("User ID Format", "UUID format or email address"),
        ("Session ID", "Must be unique integer or UUID"),
        ("Timestamp", "ISO 8601 format (YYYY-MM-DDTHH:MM:SSZ)"),
        ("File Upload", "Maximum 500MB; supported formats: Parquet, CSV, JSON"),
        ("API Key", "Must be valid JWT token with expiration < 24 hours")
    ]
    
    for rule_name, rule_details in validation_rules:
        p = doc.add_paragraph()
        p.add_run(rule_name + ": ").bold = True
        p.add_run(rule_details)
    
    add_heading_2(doc, "10.3 Privacy & Compliance")
    
    compliance_items = [
        ("Data Isolation", "Strict row-level security per client; no cross-client data access possible"),
        ("Encryption", "All data encrypted in transit (TLS 1.2+) and at rest (AES-256)"),
        ("Audit Trail", "Immutable logs of all data access with user_id, timestamp, action, resource"),
        ("Data Retention", "Session data retained for 30 days minimum, configurable per client"),
        ("GDPR Compliance", "Right to deletion implemented; PII handling adheres to GDPR requirements"),
        ("Data Masking", "Sensitive data (SSNs, credit cards) masked in logs and responses"),
        ("Access Control", "Role-based access control (RBAC) with admin, analyst, viewer roles"),
        ("API Key Rotation", "API keys expire every 90 days; rotation enforced")
    ]
    
    for item_name, item_details in compliance_items:
        p = doc.add_paragraph()
        p.add_run(item_name + ": ").bold = True
        p.add_run(item_details)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 11: API REQUIREMENTS
    # ============================================================
    add_heading_1(doc, "11. API Requirements")
    
    add_heading_2(doc, "11.1 Core API Endpoints")
    
    endpoints = [
        {
            "method": "POST",
            "path": "/api/query",
            "description": "Process natural language query",
            "params": "client_id, query, conversation_context (optional)",
            "response": "answer, context_retrieved, query_id, timestamp"
        },
        {
            "method": "GET",
            "path": "/api/clients",
            "description": "List available client profiles",
            "params": "None",
            "response": "clients (list), total (count)"
        },
        {
            "method": "GET",
            "path": "/api/create_session",
            "description": "Create new user session",
            "params": "user_id",
            "response": "session_id, user_id, created_date"
        },
        {
            "method": "POST",
            "path": "/api/fetch_attribute/dtype/v1.0",
            "description": "Upload and analyze file",
            "params": "file (multipart), user_id, session_id",
            "response": "data_types, columns, row_count, file_size"
        },
        {
            "method": "POST",
            "path": "/api/web-search",
            "description": "Augment response with web search",
            "params": "query, num_results (optional)",
            "response": "search_results (list)"
        },
        {
            "method": "GET",
            "path": "/api/health",
            "description": "System health check",
            "params": "None",
            "response": "status, uptime, version, services_status"
        }
    ]
    
    api_table = add_table(doc, len(endpoints) + 1, 5)
    api_table.style = 'Light Grid Accent 1'
    header = api_table.rows[0].cells
    header[0].text = "Method"
    header[1].text = "Endpoint"
    header[2].text = "Description"
    header[3].text = "Parameters"
    header[4].text = "Response"
    
    for endpoint in endpoints:
        row = api_table.add_row().cells
        row[0].text = endpoint["method"]
        row[1].text = endpoint["path"]
        row[2].text = endpoint["description"]
        row[3].text = endpoint["params"]
        row[4].text = endpoint["response"]
    
    add_heading_2(doc, "11.2 Standard Response Format")
    
    add_normal_paragraph(doc, "All API responses follow standardized JSON structure:")
    
    response_example = doc.add_paragraph()
    response_example.add_run("{\n"
        "  \"status\": \"success|error\",\n"
        "  \"query_id\": \"unique-uuid\",\n"
        "  \"timestamp\": \"2024-01-15T10:30:00Z\",\n"
        "  \"data\": { /* response payload */ },\n"
        "  \"metadata\": {\n"
        "    \"execution_time_ms\": 1234,\n"
        "    \"client_id\": \"plambo\",\n"
        "    \"version\": \"1.0\"\n"
        "  },\n"
        "  \"error\": null  // populated only on error\n"
        "}")
    response_example.style = 'Intense Quote'
    
    add_heading_2(doc, "11.3 Authentication & Authorization")
    
    auth_items = [
        ("Bearer Token", "All endpoints require Authorization: Bearer <JWT_TOKEN> header"),
        ("Token Expiry", "JWT tokens expire in 24 hours; refresh token mechanism required"),
        ("Scope-Based Access", "Tokens include client_id scope; API validates authorization"),
        ("Rate Limiting", "1000 requests/hour per authenticated user; returns 429 Too Many Requests"),
        ("CORS", "CORS enabled for whitelisted frontend domains only")
    ]
    
    for auth_type, auth_desc in auth_items:
        p = doc.add_paragraph()
        p.add_run(auth_type + ": ").bold = True
        p.add_run(auth_desc)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 12: ACCEPTANCE CRITERIA SUMMARY
    # ============================================================
    add_heading_1(doc, "12. Acceptance Criteria Summary")
    
    add_heading_2(doc, "12.1 Functional Acceptance")
    
    func_acceptance = [
        ("Query Processing", "Queries processed accurately with context awareness; responses < 5s (95%)"),
        ("Multi-Client Isolation", "Zero cross-client data access; authorization validated on every operation"),
        ("Session Management", "Sessions persist 30 days minimum; context preserved across interactions"),
        ("File Processing", "Files processed and indexed within 30 seconds; data types detected accurately"),
        ("API Functionality", "All documented endpoints functional and responding with correct formats"),
        ("Error Handling", "Invalid queries return descriptive errors; graceful fallback on service failure")
    ]
    
    for acceptance_type, acceptance_criteria in func_acceptance:
        p = doc.add_paragraph()
        p.add_run(acceptance_type + ": ").bold = True
        p.add_run(acceptance_criteria)
    
    add_heading_2(doc, "12.2 Non-Functional Acceptance")
    
    nfunc_acceptance = [
        ("Performance", "95th percentile response time < 5 seconds; throughput ≥ 100 req/s"),
        ("Availability", "99.5% uptime SLA maintained over 30-day period"),
        ("Security", "All authentication/authorization tests pass; no security vulnerabilities in audit"),
        ("Scalability", "Horizontal scaling supports additional instances without code changes"),
        ("Maintainability", "Code coverage ≥ 80%; all endpoints documented via OpenAPI")
    ]
    
    for acceptance_type, acceptance_criteria in nfunc_acceptance:
        p = doc.add_paragraph()
        p.add_run(acceptance_type + ": ").bold = True
        p.add_run(acceptance_criteria)
    
    add_heading_2(doc, "12.3 User Acceptance")
    
    user_acceptance = [
        ("Business Value", "Analysts can answer business questions 10x faster than manual approach"),
        ("Ease of Use", "New users complete first query without documentation (trainable in < 5 minutes)"),
        ("Reliability", "Users experience no unexpected errors in 95% of interactions"),
        ("Performance", "Response time acceptable for 90% of user interactions (< 5 seconds perceived wait)")
    ]
    
    for acceptance_type, acceptance_criteria in user_acceptance:
        p = doc.add_paragraph()
        p.add_run(acceptance_type + ": ").bold = True
        p.add_run(acceptance_criteria)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 13: RELEASE PLAN & ROADMAP
    # ============================================================
    add_heading_1(doc, "13. Release Plan & Roadmap")
    
    add_heading_2(doc, "13.1 Version 1.0 (Current - Target: Q1 2025)")
    
    v1_features = [
        "✓ Core query processing with vector store retrieval",
        "✓ Multi-client profile management",
        "✓ Session management with conversation history",
        "✓ REST API endpoints (/api/query, /api/clients, /api/create_session)",
        "✓ LLM integration (Ollama local)",
        "✓ Error handling and validation",
        "✓ Health monitoring endpoint"
    ]
    
    for feature in v1_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_heading_2(doc, "13.2 Version 1.1 (Target: Q2 2025)")
    
    v11_features = [
        "Web search integration for augmented responses",
        "Advanced file processing and data type detection",
        "Improved LLM model support (Gemini, Claude integration)",
        "Session export functionality",
        "Performance optimization for large vector stores",
        "Admin dashboard for system monitoring",
        "Rate limiting and quota management"
    ]
    
    for feature in v11_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_heading_2(doc, "13.3 Version 1.2+ (Target: Q3-Q4 2025)")
    
    v12_features = [
        "Frontend web application (currently backend only)",
        "Advanced analytics and custom report generation",
        "Slack/Teams integration for query submission",
        "Mobile application (iOS/Android)",
        "Real-time data streaming support",
        "Graph database integration (Neo4j)",
        "Fine-tuned domain-specific LLM models"
    ]
    
    for feature in v12_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_heading_2(doc, "13.4 Dependencies & Milestones")
    
    milestone_table = add_table(doc, 1, 3)
    milestone_table.style = 'Light Grid Accent 1'
    header = milestone_table.rows[0].cells
    header[0].text = "Milestone"
    header[1].text = "Target Date"
    header[2].text = "Deliverables"
    
    milestones = [
        ("Core API Launch", "March 2025", "v1.0 stable API, documentation, test coverage ≥ 80%"),
        ("First Client Deployment", "April 2025", "Production deployment with one enterprise client"),
        ("Feature Release 1.1", "June 2025", "Web search, advanced analytics, admin dashboard"),
        ("Multi-Client Scale", "August 2025", "5+ clients deployed, performance optimized"),
        ("Frontend Launch", "October 2025", "Web UI available, user-facing features complete")
    ]
    
    for milestone_name, date, deliverables in milestones:
        row = milestone_table.add_row().cells
        row[0].text = milestone_name
        row[1].text = date
        row[2].text = deliverables
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 14: RISKS & MITIGATION
    # ============================================================
    add_heading_1(doc, "14. Risks & Mitigation")
    
    add_heading_2(doc, "14.1 Technical Risks")
    
    risks = [
        {
            "risk": "LLM Service Unavailability",
            "impact": "HIGH - Queries cannot be processed",
            "mitigation": "Implement local Ollama fallback; circuit breaker pattern for API calls; graceful degradation"
        },
        {
            "risk": "Vector Store Scalability",
            "impact": "MEDIUM - Performance degradation with >1M documents",
            "mitigation": "Partition FAISS indices by time/category; implement hierarchical search; cache frequently accessed vectors"
        },
        {
            "risk": "Cross-Client Data Breach",
            "impact": "CRITICAL - Compliance violation, data leak",
            "mitigation": "Row-level security at DB layer; SQL injection prevention; regular penetration testing; audit logging"
        },
        {
            "risk": "Performance Degradation Under Load",
            "impact": "MEDIUM - SLA violation, user experience impact",
            "mitigation": "Load testing, horizontal scaling, database query optimization, caching strategies"
        },
        {
            "risk": "LLM Hallucinations",
            "impact": "MEDIUM - Incorrect insights provided to users",
            "mitigation": "System prompts with constraints; source document citations; confidence scoring; user feedback loop"
        }
    ]
    
    for risk_item in risks:
        add_heading_3(doc, risk_item["risk"])
        
        p = doc.add_paragraph()
        p.add_run("Impact: ").bold = True
        p.add_run(risk_item["impact"])
        
        p = doc.add_paragraph()
        p.add_run("Mitigation: ").bold = True
        p.add_run(risk_item["mitigation"])
    
    add_heading_2(doc, "14.2 Organizational Risks")
    
    org_risks = [
        {
            "risk": "Key Person Dependency",
            "mitigation": "Documentation; knowledge transfer sessions; cross-training team members"
        },
        {
            "risk": "Scope Creep",
            "mitigation": "Strict feature acceptance criteria; versioned releases; prioritization process"
        },
        {
            "risk": "Budget/Resource Constraints",
            "mitigation": "Phased rollout; prioritize MVP features; optimize resource allocation"
        }
    ]
    
    for risk_item in org_risks:
        p = doc.add_paragraph()
        p.add_run(risk_item["risk"] + ": ").bold = True
        p.add_run(risk_item["mitigation"])
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 15: APPENDICES
    # ============================================================
    add_heading_1(doc, "15. Appendices")
    
    add_heading_2(doc, "15.1 Glossary")
    
    glossary_items = [
        ("Conversation History", "Record of previous questions and answers in a session"),
        ("Client/Profile", "Organizational entity with isolated data and configuration"),
        ("Embedding Vector", "Numerical representation of text for semantic similarity computation"),
        ("Fallback Response", "Pre-defined response when primary processing fails"),
        ("FAISS", "Facebook AI Similarity Search - vector database for semantic retrieval"),
        ("Multi-Tenant", "Single application instance serving multiple independent organizations"),
        ("RAG", "Retrieval-Augmented Generation - combining retrieval with LLM generation"),
        ("Vector Store", "Database storing embedding vectors for similarity search"),
        ("Session", "Stateful conversation context maintained across multiple interactions")
    ]
    
    for term, definition in glossary_items:
        p = doc.add_paragraph()
        p.add_run(term + ": ").bold = True
        p.add_run(definition)
    
    add_heading_2(doc, "15.2 External References")
    
    refs = [
        "Flask Documentation: https://flask.palletsprojects.com/",
        "FAISS GitHub: https://github.com/facebookresearch/faiss",
        "Ollama: https://ollama.ai/",
        "Google Gemini API: https://ai.google.dev/",
        "ISO/IEC/IEEE 29148:2018 Standard",
        "OpenAPI 3.0 Specification",
        "OWASP API Security Guidelines"
    ]
    
    for ref in refs:
        doc.add_paragraph(ref, style='List Bullet')
    
    add_heading_2(doc, "15.3 Document Information")
    
    doc_info_table = add_table(doc, 1, 2)
    doc_info_table.style = 'Light Grid Accent 1'
    header = doc_info_table.rows[0].cells
    header[0].text = "Field"
    header[1].text = "Value"
    
    doc_info = [
        ("Document Title", "Product Requirements Document - Plambo Backend (TatvaAI)"),
        ("Version", "1.0"),
        ("Date Created", datetime.now().strftime("%B %d, %Y")),
        ("Last Updated", datetime.now().strftime("%B %d, %Y")),
        ("Status", "Draft"),
        ("Classification", "Internal - Confidential"),
        ("Project", "Plambo Conversational BI Platform"),
        ("Owner", "Product Management"),
        ("Stakeholders", "Engineering, Product, Architecture, Security, Operations")
    ]
    
    for field, value in doc_info:
        row = doc_info_table.add_row().cells
        row[0].text = field
        row[1].text = value
    
    add_heading_2(doc, "15.4 Version History")
    
    version_table = add_table(doc, 1, 4)
    version_table.style = 'Light Grid Accent 1'
    header = version_table.rows[0].cells
    header[0].text = "Version"
    header[1].text = "Date"
    header[2].text = "Author"
    header[3].text = "Changes"
    
    version_history = [
        ("1.0", datetime.now().strftime("%B %d, %Y"), "Product Team", "Initial PRD creation based on system analysis")
    ]
    
    for version, date, author, changes in version_history:
        row = version_table.add_row().cells
        row[0].text = version
        row[1].text = date
        row[2].text = author
        row[3].text = changes
    
    # Save document
    output_path = r"c:\9Trix\CODE\Plambo\plamboBE\Plambo_PRD_v1.0.docx"
    doc.save(output_path)
    print(f"PRD document saved successfully to: {output_path}")
    return output_path

if __name__ == "__main__":
    output_file = create_prd()
    print(f"\n✓ Document generation complete: {output_file}")

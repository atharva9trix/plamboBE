#!/usr/bin/env python3
"""
Incremental PRD Generator - Sections 7-8
Non-Functional Requirements & System Architecture
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class PlamboPRDSections78:
    """Generator for Sections 7-8 of PRD"""
    
    def __init__(self, existing_doc_path):
        """Load existing document"""
        self.doc = Document(existing_doc_path)
        print(f"✓ Loaded existing document: {existing_doc_path}")
    
    def add_heading_1(self, text):
        """Add Heading 1"""
        heading = self.doc.add_heading(text, level=1)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(12)
        heading_format.space_after = Pt(6)
        return heading
    
    def add_heading_2(self, text):
        """Add Heading 2"""
        heading = self.doc.add_heading(text, level=2)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(10)
        heading_format.space_after = Pt(4)
        return heading
    
    def add_heading_3(self, text):
        """Add Heading 3"""
        heading = self.doc.add_heading(text, level=3)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(8)
        heading_format.space_after = Pt(2)
        return heading
    
    def add_paragraph(self, text, bold=False, style=None):
        """Add paragraph"""
        para = self.doc.add_paragraph(text, style=style)
        para_format = para.paragraph_format
        para_format.space_before = Pt(0)
        para_format.space_after = Pt(6)
        para_format.line_spacing = 1.15
        if bold:
            for run in para.runs:
                run.bold = True
        return para
    
    def add_bullet_list(self, items):
        """Add bulleted list"""
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')
            para_format = para.paragraph_format
            para_format.space_after = Pt(3)
            para_format.line_spacing = 1.15
    
    def add_table(self, rows, cols, headers=None):
        """Add formatted table"""
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'
        
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tcPr = header_cells[i]._element.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:shd')
                tcVAlign.set(qn('w:fill'), '4472C4')
                tcPr.append(tcVAlign)
        
        return table
    
    def generate_section_7_nfr(self):
        """Generate Section 7: Non-Functional Requirements"""
        self.add_heading_1("7. Non-Functional Requirements (NFR)")
        
        self.add_paragraph(
            "Non-Functional Requirements define the quality attributes and system constraints that ensure Plambo "
            "delivers reliable, performant, secure, and maintainable service. NFRs are organized by quality characteristic "
            "following ISO/IEC/IEEE 25010 (Product Quality)."
        )
        
        # NFR Category 1: Performance
        self.add_heading_2("7.1 Performance Requirements (NFR-P)")
        
        perf_reqs = [
            {
                "id": "NFR-P.1",
                "title": "Query Response Time",
                "requirement": "System SHALL return query responses within 5 seconds (p95 latency) for 95% of queries",
                "rationale": "Sub-5-second response enables real-time decision making; faster than traditional report generation",
                "spec": [
                    "Target p50 latency (median): 2 seconds",
                    "Target p95 latency: 5 seconds",
                    "Target p99 latency: 10 seconds",
                    "Simple queries (<10 dimensions): <1.5 seconds",
                    "Complex queries (aggregations, joins): 2-7 seconds",
                    "Forecasting queries: up to 10 seconds"
                ]
            },
            {
                "id": "NFR-P.2",
                "title": "Vector Similarity Search Latency",
                "requirement": "System SHALL retrieve top-K similar documents in <100ms using FAISS index",
                "rationale": "Fast retrieval essential for <2s overall response time",
                "spec": [
                    "FAISS search <100ms for K=5 (default)",
                    "FAISS search <150ms for K=500 (maximum)",
                    "In-memory index for <10GB corpus",
                    "Approximate search (IVF) for larger indices",
                    "Single-threaded performance optimized"
                ]
            },
            {
                "id": "NFR-P.3",
                "title": "Throughput Capacity",
                "requirement": "System SHALL handle 100+ queries per second at peak load",
                "rationale": "Enterprise-grade capacity for large organizations",
                "spec": [
                    "Sustainable: 50 queries/second continuous",
                    "Peak burst: 100 queries/second for 5 minutes",
                    "Per-profile capacity: 10 queries/second minimum",
                    "Connection pooling: 50-100 concurrent connections",
                    "Memory per query context: <100MB"
                ]
            },
            {
                "id": "NFR-P.4",
                "title": "LLM API Call Optimization",
                "requirement": "System SHALL minimize LLM API calls through intelligent caching and batching",
                "rationale": "Reduces API costs and latency",
                "spec": [
                    "Cache semantic similarity threshold: 0.95",
                    "Cache TTL: configurable (default 24 hours)",
                    "Batch similar queries (up to 5) into single LLM call",
                    "Response compression: gzip for >1KB responses",
                    "Track API call budget and alert on 80% usage"
                ]
            },
            {
                "id": "NFR-P.5",
                "title": "Database Query Performance",
                "requirement": "System SHALL execute analytical queries on DuckDB within 30 seconds",
                "rationale": "Timeout prevents resource exhaustion",
                "spec": [
                    "Query timeout: 30 seconds (configurable)",
                    "Result set size: up to 1M rows in memory",
                    "Sampling support: configurable down-sampling for large results",
                    "Query plan analysis and optimization hints",
                    "Index utilization tracking and recommendations"
                ]
            }
        ]
        
        for req in perf_reqs:
            self.add_heading_3(f"{req['id']}: {req['title']}")
            self.add_paragraph(f"Requirement: {req['requirement']}", bold=True)
            self.add_paragraph(f"Rationale: {req['rationale']}")
            self.add_paragraph("Specification:", bold=True)
            self.add_bullet_list(req['spec'])
            self.add_paragraph("")
        
        # NFR Category 2: Reliability & Availability
        self.add_heading_2("7.2 Reliability & Availability (NFR-R)")
        
        rel_reqs = [
            {
                "id": "NFR-R.1",
                "title": "System Availability (Uptime)",
                "requirement": "System SHALL maintain 99.95% availability measured on monthly basis",
                "rationale": "Enterprise SLA requirement; 22 minutes downtime/month acceptable",
                "spec": [
                    "Monthly uptime target: 99.95% (99.9% acceptable)",
                    "Measured via Availability = (Total Time - Downtime) / Total Time",
                    "Excluded downtime: Scheduled maintenance (max 1hr/month), Client-side issues",
                    "On-call support for P1 incidents (platform unavailable)",
                    "RTO (Recovery Time Objective): 15 minutes for critical failures",
                    "RPO (Recovery Point Objective): 5 minutes (data loss tolerance)"
                ]
            },
            {
                "id": "NFR-R.2",
                "title": "Error Rate & Fault Tolerance",
                "requirement": "System SHALL maintain critical error rate <0.1% and implement graceful degradation",
                "rationale": "Exceeding 0.1% impacts user trust and business decisions",
                "spec": [
                    "Critical error rate: <0.1% (1 error per 1000 queries)",
                    "Transient error retry: automatic 3 retries with exponential backoff",
                    "Partial failure handling: return partial results with quality flags",
                    "LLM API failure: fall back to raw query results without synthesis",
                    "Vector store outage: fall back to keyword search",
                    "Database outage: circuit breaker and fail-fast response"
                ]
            },
            {
                "id": "NFR-R.3",
                "title": "Data Durability & Persistence",
                "requirement": "System SHALL persist all session data and audit logs with 99.999% durability (5 nines)",
                "rationale": "Data loss audit liability and compliance violation",
                "spec": [
                    "Session data: replicated across 2+ storage backends",
                    "Audit logs: write-once, immutable storage with retention 30 years",
                    "Backup frequency: hourly incremental, daily full backups",
                    "Backup validation: automated restore testing monthly",
                    "PITR (Point-in-Time Recovery): 30-day retention minimum",
                    "Disaster recovery: failover to alternate region within 30 minutes"
                ]
            },
            {
                "id": "NFR-R.4",
                "title": "Scheduled Maintenance Window",
                "requirement": "System SHALL support scheduled maintenance with zero/minimal user impact",
                "rationale": "Enables system updates without availability violations",
                "spec": [
                    "Maintenance window: weekly (e.g., Sunday 2-3 AM UTC)",
                    "Rolling updates: no service interruption for running queries",
                    "Graceful shutdown: existing connections close after timeout",
                    "Maintenance mode: return 503 Service Unavailable during updates",
                    "Notification: 48-hour advance notice to users"
                ]
            }
        ]
        
        for req in rel_reqs:
            self.add_heading_3(f"{req['id']}: {req['title']}")
            self.add_paragraph(f"Requirement: {req['requirement']}", bold=True)
            self.add_paragraph(f"Rationale: {req['rationale']}")
            self.add_paragraph("Specification:", bold=True)
            self.add_bullet_list(req['spec'])
            self.add_paragraph("")
        
        # NFR Category 3: Security
        self.add_heading_2("7.3 Security Requirements (NFR-S)")
        
        sec_reqs = [
            {
                "id": "NFR-S.1",
                "title": "Authentication & Authorization",
                "requirement": "System SHALL enforce JWT-based authentication and role-based access control (RBAC)",
                "rationale": "Prevents unauthorized access and data exposure",
                "spec": [
                    "Authentication: JWT tokens valid for 1 hour; refresh tokens 30 days",
                    "Token validation: signature verification + expiration check on every request",
                    "RBAC enforcement: role checked at endpoint level via middleware",
                    "Session binding: tokens tied to specific session and user",
                    "MFA support: optional multi-factor authentication for admin roles",
                    "SSO integration: OIDC/SAML support for enterprise deployments"
                ]
            },
            {
                "id": "NFR-S.2",
                "title": "Data Encryption",
                "requirement": "System SHALL encrypt all sensitive data in transit and at rest using AES-256",
                "rationale": "Protects against data interception and unauthorized access",
                "spec": [
                    "In transit: HTTPS/TLS 1.3 mandatory for all connections",
                    "At rest: AES-256-GCM encryption for PII and sensitive fields",
                    "Key management: centralized KMS with automatic key rotation (annual)",
                    "Database encryption: transparent encryption at storage layer",
                    "Logging: encrypt logs containing sensitive data",
                    "Certificate pinning: optional for mobile/embedded clients"
                ]
            },
            {
                "id": "NFR-S.3",
                "title": "Input Validation & Injection Prevention",
                "requirement": "System SHALL validate all inputs and prevent SQL injection, prompt injection, and other attacks",
                "rationale": "Prevents exploitation of input parsing vulnerabilities",
                "spec": [
                    "Input validation: whitelist allowed characters for each field",
                    "SQL injection: parameterized queries exclusively; no string concatenation",
                    "Prompt injection: sanitize user input before LLM consumption",
                    "Command injection: no shell command execution with untrusted input",
                    "Path traversal: restrict file operations to designated directories",
                    "Rate limiting: 100 requests/minute per user; 1000/minute per IP"
                ]
            },
            {
                "id": "NFR-S.4",
                "title": "Audit Logging & Compliance",
                "requirement": "System SHALL maintain comprehensive audit logs for all data access and configuration changes",
                "rationale": "Enables forensic investigation and regulatory compliance",
                "spec": [
                    "Audit log contents: user, timestamp, action, resource, result, IP address",
                    "Immutable logging: append-only with tamper detection",
                    "Log retention: 30 years for regulatory compliance",
                    "PII masking: automatic redaction of sensitive data in logs",
                    "Real-time alerting: anomalous access patterns trigger alerts",
                    "Access review: monthly automated review of privileged operations"
                ]
            },
            {
                "id": "NFR-S.5",
                "title": "Compliance & Standards",
                "requirement": "System SHALL comply with GDPR, HIPAA, SOC 2 Type II, and FedRAMP standards",
                "rationale": "Enables enterprise and regulated industry deployments",
                "spec": [
                    "GDPR: user consent tracking, right-to-be-forgotten implementation",
                    "HIPAA: PHI encryption, access controls, breach notification",
                    "SOC 2: completed audit with Type II report available",
                    "FedRAMP: ATO in progress (target Q3 2025)",
                    "Data residency: support for on-premises and region-specific deployment",
                    "Vulnerability scanning: automated SAST/DAST weekly"
                ]
            }
        ]
        
        for req in sec_reqs:
            self.add_heading_3(f"{req['id']}: {req['title']}")
            self.add_paragraph(f"Requirement: {req['requirement']}", bold=True)
            self.add_paragraph(f"Rationale: {req['rationale']}")
            self.add_paragraph("Specification:", bold=True)
            self.add_bullet_list(req['spec'])
            self.add_paragraph("")
        
        # NFR Category 4: Scalability
        self.add_heading_2("7.4 Scalability & Elasticity (NFR-SC)")
        
        scale_reqs = [
            {
                "id": "NFR-SC.1",
                "title": "Horizontal Scalability",
                "requirement": "System architecture SHALL support horizontal scaling (add more servers) without code changes",
                "rationale": "Enables cost-effective capacity growth",
                "spec": [
                    "Stateless API design: no server-side session state",
                    "Load balancing: round-robin across N instances",
                    "Database connection pooling: shared pool across instances",
                    "Vector store replication: read-only replicas for load distribution",
                    "Session store: distributed cache (Redis) for session data",
                    "Scaling target: 100 queries/sec per 8-core instance"
                ]
            },
            {
                "id": "NFR-SC.2",
                "title": "Vertical Scalability",
                "requirement": "System SHALL support vertical scaling (larger servers) for single-instance deployments",
                "rationale": "Enables incremental capacity growth within organizational constraints",
                "spec": [
                    "Memory efficiency: <500MB per idle connection",
                    "CPU efficiency: linear scaling up to 64 cores",
                    "Disk I/O: support fast SSDs (>10K IOPS)",
                    "Network: gigabit or better connectivity",
                    "Tested on: 8 core to 64 core instances production-verified"
                ]
            },
            {
                "id": "NFR-SC.3",
                "title": "Data Volume Scalability",
                "requirement": "System SHALL handle growing data volumes without performance degradation",
                "rationale": "Organizations accumulate hundreds of GB of data over time",
                "spec": [
                    "Query dataset size: support up to 100GB indexed data per profile",
                    "Query result size: 1M rows in memory; streaming for larger",
                    "FAISS index scaling: partitioned indices for >10GB",
                    "Session history: support 100k concurrent sessions",
                    "Conversation history: unlimited per session (30-year retention)",
                    "Incremental indexing: add documents without reindexing"
                ]
            }
        ]
        
        for req in scale_reqs:
            self.add_heading_3(f"{req['id']}: {req['title']}")
            self.add_paragraph(f"Requirement: {req['requirement']}", bold=True)
            self.add_paragraph(f"Rationale: {req['rationale']}")
            self.add_paragraph("Specification:", bold=True)
            self.add_bullet_list(req['spec'])
            self.add_paragraph("")
        
        # NFR Category 5: Usability & Compatibility
        self.add_heading_2("7.5 Usability & Compatibility (NFR-U)")
        
        usab_reqs = [
            {
                "id": "NFR-U.1",
                "title": "API Usability & Documentation",
                "requirement": "System SHALL provide comprehensive OpenAPI 3.1 specification and interactive documentation",
                "rationale": "Reduces integration effort and support burden",
                "spec": [
                    "OpenAPI spec: complete, machine-readable, auto-generated",
                    "Interactive docs: Swagger UI for all endpoints",
                    "Code examples: Python, JavaScript, cURL samples for each operation",
                    "SDKs: Official Python and JavaScript client libraries",
                    "Error responses: consistent format with helpful error messages",
                    "API versioning: semantic versioning with deprecation window (6 months)"
                ]
            },
            {
                "id": "NFR-U.2",
                "title": "Language & Localization Support",
                "requirement": "System SHALL support natural language queries in English, Spanish, French, German provisionally",
                "rationale": "Enables global enterprise adoption",
                "spec": [
                    "Query languages: EN (MVP), ES, FR, DE (Q2 2025)",
                    "UI text: translations for common messages",
                    "Date formats: locale-aware formatting",
                    "Number formats: locale-aware currency, decimal separators",
                    "Timezone support: UTC storage, local display per user"
                ]
            },
            {
                "id": "NFR-U.3",
                "title": "Browser & Client Compatibility",
                "requirement": "System REST API SHALL be compatible with any HTTP client; no browser-specific features",
                "rationale": "Enables diverse client implementations",
                "spec": [
                    "HTTP/2 support: full HTTP/2 protocol compliance",
                    "CORS: configurable cross-origin support",
                    "Mobile clients: responsive REST API, minimal payload",
                    "Embedded clients: no external dependencies (pure HTTP)",
                    "Legacy clients: backward compatibility with HTTP/1.1"
                ]
            }
        ]
        
        for req in usab_reqs:
            self.add_heading_3(f"{req['id']}: {req['title']}")
            self.add_paragraph(f"Requirement: {req['requirement']}", bold=True)
            self.add_paragraph(f"Rationale: {req['rationale']}")
            self.add_paragraph("Specification:", bold=True)
            self.add_bullet_list(req['spec'])
            self.add_paragraph("")
        
        # NFR Category 6: Maintainability
        self.add_heading_2("7.6 Maintainability & Operations (NFR-M)")
        
        maint_reqs = [
            {
                "id": "NFR-M.1",
                "title": "Code Quality & Testing",
                "requirement": "System codebase SHALL maintain >80% test coverage and pass automated quality checks",
                "rationale": "High code quality reduces bugs and maintenance costs",
                "spec": [
                    "Unit test coverage: >80% of production code",
                    "Integration tests: >60% of API endpoints",
                    "Code style: PEP 8 compliance (Python), ESLint (JavaScript)",
                    "Static analysis: Bandit, SAST tools for security",
                    "Dependency scanning: automated vulnerability detection",
                    "Pre-commit hooks: enforce code quality on every commit"
                ]
            },
            {
                "id": "NFR-M.2",
                "title": "Observability & Monitoring",
                "requirement": "System SHALL provide comprehensive logs, metrics, and traces for operational visibility",
                "rationale": "Enables rapid incident investigation and performance optimization",
                "spec": [
                    "Logs: structured JSON format, sent to central logging (ELK/Splunk)",
                    "Metrics: Prometheus-compatible metrics exports",
                    "Traces: distributed tracing with OpenTelemetry (Jaeger backend)",
                    "Dashboards: Grafana dashboards for system health",
                    "Alerting: PagerDuty integration for critical alerts",
                    "SLO tracking: monitor against defined SLIs"
                ]
            },
            {
                "id": "NFR-M.3",
                "title": "Deployment & Release Management",
                "requirement": "System SHALL support blue-green deployments with zero downtime",
                "rationale": "Minimizes customer impact during releases",
                "spec": [
                    "Container support: Docker images for all components",
                    "Orchestration: Kubernetes manifests provided",
                    "CI/CD pipeline: automated testing and deployment",
                    "Deployment strategy: blue-green with health checks",
                    "Rollback capability: automatic rollback on health check failure",
                    "Release frequency: weekly minor releases, critical hotfixes within 24hrs"
                ]
            }
        ]
        
        for req in maint_reqs:
            self.add_heading_3(f"{req['id']}: {req['title']}")
            self.add_paragraph(f"Requirement: {req['requirement']}", bold=True)
            self.add_paragraph(f"Rationale: {req['rationale']}")
            self.add_paragraph("Specification:", bold=True)
            self.add_bullet_list(req['spec'])
            self.add_paragraph("")
        
        self.doc.add_page_break()
    
    def generate_section_8_architecture(self):
        """Generate Section 8: System Architecture Overview"""
        self.add_heading_1("8. System Architecture Overview")
        
        self.add_heading_2("8.1 Architectural Principles")
        
        self.add_paragraph(
            "The Plambo Backend is architected following these guiding principles to ensure reliability, scalability, "
            "and maintainability:"
        )
        
        principles = [
            ("Separation of Concerns", "Distinct layers (API, Service, Data) with clear interfaces"),
            ("Statelessness", "API instances are stateless; state stored in distributed cache/database"),
            ("Resilience", "Graceful degradation; fallback strategies for component failures"),
            ("Security by Design", "Encryption, authentication, and authorization at every layer"),
            ("Observability", "Comprehensive logging, metrics, and tracing throughout"),
            ("Modularity", "Loose coupling enables independent scaling and updates"),
            ("Asynchronicity", "Non-blocking I/O for high concurrency"),
            ("Cloud-Native", "Container-based, auto-scaling, managed services ready")
        ]
        
        for principle, description in principles:
            self.add_paragraph(f"{principle}: {description}")
        
        self.add_heading_2("8.2 High-Level Architecture Diagram")
        
        self.add_paragraph(
            "The Plambo Backend follows a microservices-oriented layered architecture with the following major components:"
        )
        
        self.add_heading_3("8.2.1 Architectural Layers")
        
        architecture_desc = """
┌─────────────────────────────────────────────────────────────────┐
│                      CLIENT APPLICATIONS                         │
│         (Frontend, Mobile, Third-Party Integrations)             │
└──────────────────────────┬──────────────────────────────────────┘
                           │ HTTPS REST API
┌──────────────────────────▼──────────────────────────────────────┐
│                    API GATEWAY / LOAD BALANCER                   │
│              (Rate Limiting, Auth, Request Routing)              │
└──────────────────────────┬──────────────────────────────────────┘
                           │
        ┌──────────────────┼──────────────────┐
        │                  │                  │
┌───────▼─────────┐  ┌────▼────────┐  ┌─────▼──────────┐
│ Query Controller│  │Session Ctrl │  │Config Controller│
│ Web Controller  │  │Health Check │  │Profile Manager  │
│ TatvaAI Handler │  │             │  │Auth Middleware  │
└────────┬────────┘  └────┬────────┘  └────┬──────────┘
         │                │                │
         └────────────────┼────────────────┘
                          │
         ┌────────────────┴─────────────────┐
         │  SERVICE LAYER (Business Logic)  │
         │                                   │
    ┌────▼─────────────────────────────────┐
    │ • Query Service (parsing, intention) │
    │ • Session Service (state mgmt)       │
    │ • RAG Service (retrieval + gen)      │
    │ • Profile Service (multi-tenancy)    │
    │ • Insight Service (LLM integration)  │
    └────┬──────────────────────────────────┘
         │
         ├──────────────┬───────────┬──────────────┐
         │              │           │              │
    ┌────▼────┐    ┌────▼──┐  ┌──▼────┐   ┌─────▼────┐
    │ DuckDB  │    │FAISS  │  │ Cache │   │Auth Srvr │
    │(Queries)│    │(Vec.  │  │(Redis)│   │ (TatvaAI)│
    │         │    │Search)│  │       │   │          │
    └────┬────┘    └────┬──┘  └──┬────┘   └─────┬────┘
         │              │        │              │
         └──────────────┼────────┼──────────────┘
                        │        │
    ┌───────────────────┤        ├──────────────────┐
    │                   │        │                  │
┌───▼──────┐  ┌────────▼─┐  ┌──▼─────┐  ┌────────▼──┐
│PostgreSQL│  │ S3/Blob  │  │ LLM API│  │  Metrics  │
│(Sessions)│  │ (Parquet)│  │(Gemini)│  │ (Prometheus)
└──────────┘  └──────────┘  └────────┘  └──────────┘
        """
        
        self.add_paragraph(architecture_desc, style='List Bullet')
        
        self.add_heading_3("8.2.2 Component Descriptions")
        
        components_table_data = [
            ("Component", "Responsibility", "Technology", "Criticality"),
            ("API Gateway", "Request routing, auth check, rate limiting", "Flask/Nginx", "Critical"),
            ("Query Service", "Parse intent, extract entities, generate SQL", "spaCy, BERT, LangChain", "Critical"),
            ("RAG Service", "Retrieve docs, augment generation", "FAISS, Gemini API", "Critical"),
            ("DuckDB", "Analytical query execution", "DuckDB", "Critical"),
            ("Session Store", "Conversation history persistence", "PostgreSQL", "Critical"),
            ("Cache Layer", "Session context, query cache", "Redis", "High"),
            ("Vector Store", "FAISS indices per profile", "FAISS (in-memory)", "Critical"),
            ("LLM Client", "Gemini API integration", "google-genai library", "High"),
            ("Auth Service", "JWT validation, RBAC", "TatvaAI", "Critical"),
            ("Logging", "Structured logs collection", "Python logging + ELK", "High"),
            ("Metrics", "System metrics export", "Prometheus client", "Medium"),
        ]
        
        comp_table = self.add_table(len(components_table_data), 4, headers=components_table_data[0])
        for row_idx in range(1, len(components_table_data)):
            comp_table.rows[row_idx].cells[0].text = components_table_data[row_idx][0]
            comp_table.rows[row_idx].cells[1].text = components_table_data[row_idx][1]
            comp_table.rows[row_idx].cells[2].text = components_table_data[row_idx][2]
            comp_table.rows[row_idx].cells[3].text = components_table_data[row_idx][3]
        
        self.add_heading_2("8.3 Data Flow Architecture")
        
        self.add_heading_3("8.3.1 Query Processing Flow")
        
        flow_steps = [
            "1. User Query Reception: Client submits natural language query via POST /api/query",
            "2. Authentication: API Gateway validates JWT token via TatvaAI service",
            "3. Rate Limiting: Check request quota (per-user, per-IP)",
            "4. Session Context Load: Retrieve conversation history from PostgreSQL + Redis cache",
            "5. Query Parsing: QueryService parses intent and extracts entities",
            "6. Semantic Search: Query embedded into vector space, FAISS retrieval of K documents",
            "7. Context Assembly: Retrieved documents + conversation history prepared for LLM",
            "8. SQL Generation: Construct parameterized DuckDB query from extracted entities",
            "9. Query Execution: Execute SQL with 30-second timeout",
            "10. Result Formatting: Transform DuckDB results to JSON with metadata",
            "11. Insight Synthesis: Call Gemini API with results + context for natural lang insights",
            "12. Response Assembly: Combine insights, data, metadata, confidence",
            "13. Session Update: Store new query-response exchange in session history",
            "14. Response Return: Return JSON response to client with 200 OK"
        ]
        
        self.add_bullet_list(flow_steps)
        
        self.add_heading_3("8.3.2 Error Handling Flow")
        
        error_flow = [
            "Invalid Input (400): Input validation fails → return error details",
            "Authentication Failure (401): JWT invalid/expired → request user re-authentication",
            "Authorization Failure (403): User lacks required role → return permission denied",
            "Query Timeout (504): SQL execution exceeds 30s → return partial results + warning",
            "LLM API Failure (500): Gemini call fails → return raw results without synthesis",
            "Vector Store Failure (500): FAISS unavailable → fall back to keyword search",
            "Database Failure (500): PostgreSQL/DuckDB down → activate circuit breaker, return 503",
            "Rate Limit Exceeded (429): User quota exceeded → queue request, return 429 with Retry-After"
        ]
        
        self.add_bullet_list(error_flow)
        
        self.add_heading_2("8.4 Deployment Architecture")
        
        self.add_heading_3("8.4.1 Deployment Models")
        
        self.add_paragraph("Supported Deployment Scenarios:")
        
        deployment_models = [
            ("Cloud-Native (Recommended)", "Kubernetes cluster on AWS/GCP/Azure with managed databases and auto-scaling"),
            ("On-Premises", "Docker containers with external PostgreSQL/Redis, manual scaling"),
            ("Hybrid", "Plambo backend on-prem, LLM calls through secure gateway to cloud"),
            ("Single-Server (Dev/Test)", "Monolithic deployment on single EC2/VM with SQLite")
        ]
        
        for model, desc in deployment_models:
            self.add_paragraph(f"{model}: {desc}")
        
        self.add_heading_3("8.4.2 Infrastructure Requirements")
        
        infra_data = [
            ("Component", "CPU", "Memory", "Storage", "Network"),
            ("API Instances (min 2)", "4 cores", "8GB", "10GB", "1Gbps"),
            ("PostgreSQL", "4 cores", "16GB", "500GB (SSD)", "1Gbps"),
             ("Redis Cache", "2 cores", "8GB", "100GB", "1Gbps"),
            ("Vector Store (FAISS)", "N/A", "32GB+", "N/A", "N/A"),
            ("Load Balancer", "2 cores", "4GB", "5GB", "10Gbps"),
            ("Total (Production)", "12-24 cores", "40-64GB", "1TB+", "10Gbps+")
        ]
        
        infra_table = self.add_table(len(infra_data), 5, headers=infra_data[0])
        for row_idx in range(1, len(infra_data)):
            for col_idx in range(5):
                infra_table.rows[row_idx].cells[col_idx].text = infra_data[row_idx][col_idx]
        
        self.add_heading_3("8.4.3 Technology Stack")
        
        tech_stack = [
            ("Layer", "Technology", "Justification"),
            ("Runtime", "Python 3.10+", "Development speed, ecosystem maturity"),
            ("Web Framework", "Flask 2.3+", "Lightweight, extensible microservices"),
            ("Data Processing", "Pandas, NumPy", "Data manipulation and analysis"),
            ("Query Execution", "DuckDB", "Analytical queries, in-process SQL"),
            ("Vector Search", "FAISS", "Semantic similarity, performance, open-source"),
            ("NLP/Embeddings", "spaCy, sentence-transformers", "Entity extraction, embeddings"),
            ("LLM", "Google Gemini 2.5", "SOTA reasoning, API availability"),
            ("Session Store", "PostgreSQL 14+", "Durability, ACID compliance"),
            ("Cache", "Redis 7.0+", "High-speed in-memory state"),
            ("Async Tasks", "Celery (optional)", "Future background job support"),
            ("API Documentation", "OpenAPI 3.1, Swagger", "Self-documenting APIs"),
            ("Testing", "pytest, pytest-cov", "Unit & integration testing"),
            ("Monitoring", "Prometheus, Grafana, ELK", "Metrics, dashboards, logs"),
            ("Containerization", "Docker 20.10+", "Reproducible deployments"),
            ("Orchestration", "Kubernetes 1.24+", "Production workload management"),
        ]
        
        tech_table = self.add_table(len(tech_stack), 3, headers=tech_stack[0])
        for row_idx in range(1, len(tech_stack)):
            tech_table.rows[row_idx].cells[0].text = tech_stack[row_idx][0]
            tech_table.rows[row_idx].cells[1].text = tech_stack[row_idx][1]
            tech_table.rows[row_idx].cells[2].text = tech_stack[row_idx][2]
        
        self.add_heading_2("8.5 Integration Points")
        
        self.add_heading_3("8.5.1 External Integrations")
        
        integrations = [
            ("Google Gemini API", "LLM-powered insight generation and query understanding", "Critical"),
            ("External Search APIs", "Web search context enrichment (future)", "Low"),
            ("LDAP/Active Directory", "Enterprise authentication and user sync", "High"),
            ("SMTP/Email", "User notifications and alerts", "Medium"),
            ("Webhook Endpoints", "Async notification delivery to clients", "Medium"),
            ("S3 / Cloud Storage", "Parquet file and backup storage", "High"),
            ("Monitoring/SIEM", "Security event export (Splunk, ELK)", "Medium"),
        ]
        
        for integration, purpose, importance in integrations:
            self.add_paragraph(f"{integration}: {purpose} ({importance} importance)")
        
        self.add_heading_3("8.5.2 API Versioning Strategy")
        
        versioning = [
            "Major versions (breaking changes): v1, v2 (separate endpoints /api/v1/*, /api/v2/*)",
            "Minor versions (backward compatible): v1.1, v1.2 (same endpoint, new optional parameters)",
            "Deprecation window: 6 months advance notice before removing APIs",
            "Documentation strategy: maintain docs for current and previous major version only",
            "Support policy: security fixes for current version, critical fixes for prior version (6 months)"
        ]
        
        self.add_bullet_list(versioning)
        
        self.doc.add_page_break()
    
    def save_document(self, filename):
        """Save document"""
        self.doc.save(filename)
        print(f"✓ Document saved: {filename}")
        return filename


def main():
    """Generate Sections 7-8"""
    print("\n" + "=" * 70)
    print("PLAMBO PRD GENERATOR - SECTIONS 7-8 EXTENSION")
    print("=" * 70)
    
    existing_file = r"c:\9Trix\CODE\Plambo\plamboBE\Plambo_PRD_v1.0_Sections_1-6.docx"
    
    print(f"\n[STEP 1] Loading existing document...")
    generator = PlamboPRDSections78(existing_file)
    
    print("[STEP 2] Generating Section 7: Non-Functional Requirements...")
    generator.generate_section_7_nfr()
    
    print("[STEP 3] Generating Section 8: System Architecture...")
    generator.generate_section_8_architecture()
    
    output_file = r"c:\9Trix\CODE\Plambo\plamboBE\Plambo_PRD_v1.0_Sections_1-8.docx"
    generator.save_document(output_file)
    
    print("\n" + "=" * 70)
    print("✓ SECTIONS 7-8 COMPLETE - DOCUMENT EXTENDED")
    print("=" * 70)
    print(f"\nOutput: Plambo_PRD_v1.0_Sections_1-8.docx")
    print("\nSections now included:")
    print("  ✓ Section 1-6: Previously generated")
    print("  ✓ Section 7: Non-Functional Requirements (20 NFRs)")
    print("  ✓ Section 8: System Architecture Overview")
    print("\n→ Ready for next sections. Type 'Continue' to proceed to Sections 9-11.")
    
    return output_file


if __name__ == "__main__":
    main()

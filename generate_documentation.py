#!/usr/bin/env python3
"""
Generate comprehensive technical documentation for plamboBE as Word document
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_documentation():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('PlamBO Backend System', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Technical Documentation')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph('Generated: February 13, 2026')
    doc.add_paragraph('Repository: plamboBE (branch: patch1.1)')
    doc.add_paragraph('Owner: atharva9trix')
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Overview',
        '2. Architecture Overview', 
        '3. Environment & Prerequisites',
        '4. File-by-File Documentation',
        '5. Code Block Deep Dive',
        '6. Functions, Classes & Methods',
        '7. State & Data Flow',
        '8. APIs / Interfaces',
        '9. Configuration & Constants',
        '10. Error Handling & Logging',
        '11. Security Considerations',
        '12. Performance & Scalability',
        '13. Limitations & Technical Debt',
        '14. Extension & Modification Guide'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # SECTION 1: Executive Overview
    doc.add_heading('SECTION 1: Executive Overview', level=1)
    
    doc.add_heading('1.1 System Purpose', level=2)
    doc.add_paragraph(
        'PlamBO Backend (plamboBE) is a comprehensive enterprise data analysis and knowledge management '
        'system designed to serve multiple business clients simultaneously. The system provides a unified '
        'platform for document indexing, semantic search, conversational query processing, data analysis, '
        'and intelligent business intelligence generation. It acts as a bridge between raw business data '
        'and actionable insights, combining vector search capabilities with large language models (LLMs) '
        'for context-aware query responses.'
    )
    
    doc.add_heading('1.2 Problem Statement', level=2)
    doc.add_paragraph(
        'Modern enterprises face multiple challenges: (1) Managing knowledge bases across multiple business '
        'units is fragmented and inefficient. (2) Finding relevant information from large document repositories '
        'requires advanced semantic understanding, not just keyword matching. (3) Converting raw data analysis '
        'queries into actionable business intelligence requires both SQL expertise and domain knowledge. '
        '(4) Multi-client deployments require complete data isolation and access control to prevent information '
        'leakage. (5) Session management across long-running analytical queries demands persistence and '
        'resumability. PlamBO Backend solves these challenges through a unified, multi-tenant architecture.'
    )
    
    doc.add_heading('1.3 Intended Users', level=2)
    doc.add_paragraph(
        'End Users - Business analysts, data scientists, and non-technical stakeholders querying knowledge bases and datasets.'
    )
    doc.add_paragraph(
        'Administrators - System operators managing deployments, configurations, and client profiles.'
    )
    doc.add_paragraph(
        'Developers - Engineers extending functionality or integrating plamboBE into larger enterprise systems.'
    )
    doc.add_paragraph(
        'Data Stewards - Teams responsible for maintaining data quality and knowledge base updates.'
    )
    
    doc.add_heading('1.4 Deployment Context', level=2)
    doc.add_paragraph(
        'PlamBO Backend is a Python-based Flask REST API designed for deployment in containerized environments '
        '(Docker) or traditional servers. It runs as a background service on port 8000 by default. The system '
        'requires PostgreSQL for transactional data, DuckDB for in-memory analytics, FAISS for vector search, '
        'and Ollama for local LLM serving. All components can run on a single server or be distributed across '
        'multiple machines. The system is designed to handle multiple concurrent client requests with proper '
        'session isolation and request validation.'
    )
    
    doc.add_heading('1.5 Key Assumptions', level=2)
    assumptions = [
        'All client knowledge bases are pre-indexed and available as FAISS vector stores.',
        'PostgreSQL database is running and accessible with provided credentials.',
        'Ollama LLM service is running locally at localhost:11434 with gemma3:1b model.',
        'CPU-only inference is acceptable (no GPU required or available).',
        'Vector stores contain sufficient documents to answer client queries.',
        'Client IDs uniquely identify knowledge profiles and access scopes.',
        'Session data is transient and stored in PostgreSQL, not persistent across restarts.',
        'File uploads are temporary and processed synchronously in request lifecycle.',
        'No sensitive data (e.g., PII, credentials) is stored in file uploads.',
        'Cross-profile queries are forbidden for security and data isolation.'
    ]
    for assumption in assumptions:
        doc.add_paragraph(assumption, style='List Bullet')
    
    doc.add_page_break()
    
    # SECTION 2: Architecture Overview
    doc.add_heading('SECTION 2: Architecture Overview', level=1)
    
    doc.add_heading('2.1 High-Level Architecture', level=2)
    doc.add_paragraph(
        'PlamBO Backend follows a layered, multi-component architecture with clear separation of concerns:'
    )
    
    doc.add_paragraph('Entry Layer', style='List Bullet')
    doc.add_paragraph(
        'Flask REST API serving HTTP endpoints. All requests hit /api/* routes registered with blueprints.',
        style='List Bullet 2'
    )
    
    doc.add_paragraph('Validation Layer', style='List Bullet')
    doc.add_paragraph(
        'QueryValidator ensures incoming payloads contain required fields and valid client IDs before processing.',
        style='List Bullet 2'
    )
    
    doc.add_paragraph('Business Logic Layer', style='List Bullet')
    doc.add_paragraph(
        'QueryService and WebService process requests, load vector stores, and coordinate operations.',
        style='List Bullet 2'
    )
    
    doc.add_paragraph('Data Retrieval Layer', style='List Bullet')
    doc.add_paragraph(
        'VectorStoreLoader queries FAISS indexes to retrieve context documents based on semantic similarity.',
        style='List Bullet 2'
    )
    
    doc.add_paragraph('LLM Processing Layer', style='List Bullet')
    doc.add_paragraph(
        'LLMProcessor calls Ollama for natural language generation with context guardrails.',
        style='List Bullet 2'
    )
    
    doc.add_paragraph('Persistence Layer', style='List Bullet')
    doc.add_paragraph(
        'PostgreSQL stores session data, file metadata, and user session information. DuckDB provides in-memory analytics for complex queries.',
        style='List Bullet 2'
    )
    
    doc.add_paragraph('File Handling Layer', style='List Bullet')
    doc.add_paragraph(
        'Readwrite handles multi-format file reading/writing; update_parquet manages Parquet conversions.',
        style='List Bullet 2'
    )
    
    doc.add_heading('2.2 Data Flow', level=2)
    doc.add_paragraph('Typical query execution flow:')
    
    steps = [
        'Client sends POST /api/query with {"client_id": "plambo", "query": "..."} JSON payload.',
        'query_controller.query() receives request, calls QueryValidator.validate to ensure client_id is valid.',
        'QueryService.process retrieves vector store for client_id via profile_manager.load_profile().',
        'VectorStoreLoader.retrieve() encodes query, searches FAISS index, returns top-k documents with scores.',
        'llm_processor.process_query() builds system prompt with context, calls Ollama API, streams response.',
        'Response is formatted and returned to client as JSON: {"status": "success", "answer": "..."}',
        'If Ollama fails or no context found, fallback response: "This question is not within the scope..."'
    ]
    
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_page_break()
    
    # SECTION 3: Environment & Prerequisites
    doc.add_heading('SECTION 3: Environment & Prerequisites', level=1)
    
    doc.add_heading('3.1 Required Software & Versions', level=2)
    software = [
        ('Python', '3.8+'),
        ('Flask', 'Latest (used for REST API)'),
        ('PostgreSQL', '11+ (transactional database)'),
        ('DuckDB', 'Latest (in-memory analytics)'),
        ('FAISS', 'Latest with CPU-only support'),
        ('Ollama', 'Latest (gemma3:1b model required)'),
        ('sentence-transformers', 'with all-MiniLM-L6-v2 model'),
        ('pandas', 'Latest (data manipulation)'),
        ('PyArrow', 'Latest (Parquet support)'),
        ('SQLAlchemy', 'Latest (ORM for PostgreSQL)')
    ]
    for sw, ver in software:
        doc.add_paragraph(f'{sw}: {ver}')
    
    doc.add_heading('3.2 Folder Structure & Purpose', level=2)
    
    folders = {
        'src/config/': 'Configuration files: config.py, settings.py, db_config.yml, config.yml',
        'src/controllers/': 'Flask blueprint endpoints for query, web, health, tatva operations',
        'src/services/': 'Business logic for query processing and web search',
        'src/profiles/': 'Profile (client) management and vector store loading',
        'src/llm/': 'LLM processor for Ollama integration',
        'src/db_connection/': 'Database engine and connection management',
        'src/validators/': 'Input validation for API payloads',
        'src/data_analysis_dckdb/': 'Analytics processing with DuckDB and Google Gemini API',
        'src/vector_stores/': 'Pre-built FAISS indexes for each client profile',
        'data/': 'Runtime directory for file uploads and parquet conversions'
    }
    
    for folder, purpose in folders.items():
        doc.add_paragraph(f'{folder}: {purpose}')
    
    doc.add_page_break()
    
    # SECTION 4: File-by-File Documentation
    doc.add_heading('SECTION 4: File-by-File Documentation', level=1)
    
    doc.add_heading('4.1 app.py - Application Bootstrap', level=2)
    doc.add_paragraph('Purpose: Main entry point of the Flask application. Initializes the Flask app and registers all blueprints.')
    doc.add_paragraph('Why it exists: To provide a clean, single-responsibility bootstrap file for application startup and configuration.')
    doc.add_paragraph('Interaction: Imports all controllers and registers them as blueprints. Calls register_error_handlers to set up global error handlers.')
    doc.add_paragraph('Execution: Run as: python app.py. Starts Flask development server on 0.0.0.0:8000.')
    
    doc.add_heading('4.2 src/config/settings.py - Application Settings', level=2)
    doc.add_paragraph('Purpose: Centralized configuration for all clients, models, and API settings.')
    doc.add_paragraph('Why it exists: To provide a single source of truth for all application constants and avoid hardcoding values throughout codebase.')
    doc.add_paragraph('Interaction: Imported by controllers, services, and profile manager to access client profiles and settings.')
    doc.add_paragraph('Key configurations: CLIENTS dict defines all multi-tenant profiles. EMBEDDING_MODEL specifies sentence-transformers model. LLM_CONFIG defines Ollama settings.')
    
    doc.add_heading('4.3 src/config/config.py - Environment Configuration Loader', level=2)
    doc.add_paragraph('Purpose: Load YAML configuration files and expose as class attributes with environment variable substitution.')
    doc.add_paragraph('Why it exists: To support multiple deployment environments (DEVELOPMENT, PRODUCTION) without code changes.')
    doc.add_paragraph('Interaction: Used by db_engine.py to get database credentials, by tatva_ai.py to get paths.')
    doc.add_paragraph('Execution: Config(environment="DEVELOPMENT") loads YAML files and substitutes ${ENV_VAR} patterns from environment.')
    
    doc.add_heading('4.4 src/controllers/query_controller.py - Query API Endpoint', level=2)
    doc.add_paragraph('Purpose: Handle POST /api/query requests for querying knowledge base profiles.')
    doc.add_paragraph('Interaction: Calls validator to check payload, invokes QueryService.process() to execute business logic.')
    doc.add_paragraph('Execution: Synchronous. Returns JSON response with status and answer.')
    doc.add_paragraph('Routes: POST /api/query for knowledge base queries; GET /api/clients to list available profiles.')
    
    doc.add_heading('4.5 src/services/query_service.py - Query Processing Logic', level=2)
    doc.add_paragraph('Purpose: Orchestrate query processing: load vector store, retrieve documents, call LLM.')
    doc.add_paragraph('Interaction: Orchestrates between ProfileManager, VectorStoreLoader, and LLMProcessor.')
    doc.add_paragraph('Execution: Synchronous request-response cycle. Loads profile, retrieves context, processes through LLM.')
    
    doc.add_heading('4.6 src/profiles/manager.py - Profile Management', level=2)
    doc.add_paragraph('Purpose: Central registry for all client profiles and their vector stores with lazy loading.')
    doc.add_paragraph('Interaction: Validates profile IDs against CLIENTS config. Caches loaded VectorStoreLoader instances.')
    doc.add_paragraph('Execution: load_profile(client_id) returns cached or newly-loaded VectorStoreLoader.')
    
    doc.add_heading('4.7 src/profiles/loader.py - Vector Store Loader', level=2)
    doc.add_paragraph('Purpose: Load FAISS index and metadata for a specific profile.')
    doc.add_paragraph('Interaction: Loads sentence-transformer model and FAISS index on first instantiation.')
    doc.add_paragraph('Execution: Lazy loading. Caches instances in ProfileManager._stores. On first call, loads model and index.')
    doc.add_paragraph('retrieve() method: Transforms query to embedding, searches FAISS, returns top-k relevant documents with similarity scores.')
    
    doc.add_heading('4.8 src/llm/processor.py - LLM Integration with Guardrails', level=2)
    doc.add_paragraph('Purpose: Process queries through Ollama with strict guardrails to prevent hallucination.')
    doc.add_paragraph('Interaction: Receives query string and retrieved documents. Builds prompt with context. Calls Ollama API.')
    doc.add_paragraph('Execution: Applies 4 guardrails: context check, context string building, prompt construction, not-found detection.')
    doc.add_paragraph('System Contract: Enforces that LLM uses only provided context; forbids general knowledge use.')
    
    doc.add_heading('4.9 src/validators/query_validator.py - Input Validation', level=2)
    doc.add_paragraph('Purpose: Validate incoming query payloads.')
    doc.add_paragraph('Checks: (1) Payload not null. (2) client_id in CLIENTS. (3) query not empty.')
    doc.add_paragraph('Execution: Raises ValueError if validation fails; caught by error handler and returned as 400 response.')
    
    doc.add_heading('4.10 src/db_connection/db_engine.py - Database Engine', level=2)
    doc.add_paragraph('Purpose: Manage PostgreSQL connections and database operations.')
    doc.add_paragraph('Engine class: Creates SQLAlchemy engine and connections. connect_engine() returns (connection, status).')
    doc.add_paragraph('Read_Write class: fetch_data() executes SELECT, post_data() inserts, update_data() executes UPDATE/DELETE.')
    doc.add_paragraph('Interaction: Used by TatvaAIMain, DataDescribe for session persistence and data operations.')
    
    doc.add_page_break()
    
    # SECTION 5: Code Block Deep Dive
    doc.add_heading('SECTION 5: Code Block Deep Dive', level=1)
    
    doc.add_heading('5.1 QueryService.process() - Query Pipeline Orchestration', level=2)
    doc.add_paragraph(
        'What it does: Orchestrates the complete query pipeline. Loads the appropriate profile, retrieves context, processes through LLM, returns structured response.'
    )
    doc.add_paragraph('Why it exists: Separates business logic from HTTP handling. Enables reuse from multiple endpoints.')
    doc.add_paragraph('Inputs: payload dict with client_id, query, conversation_context.')
    doc.add_paragraph('Outputs: dict with status=success, answer, context_retrieved count.')
    doc.add_paragraph('Side effects: Loads vector store into memory (cached). Makes external HTTP call to Ollama. May timeout.')
    doc.add_paragraph('What breaks if removed: Query processing becomes impossible; all /api/query calls fail.')
    doc.add_paragraph('Edge cases handled: (1) Empty documents: Falls back to standard message. (2) LLM timeout: Returns error message. (3) Invalid client_id: Validator catches upstream.')
    doc.add_paragraph('Edge cases NOT handled: (1) Malformed query payload crashes without validation; caught by error handler. (2) Crashed Ollama service does not retry.')
    
    doc.add_heading('5.2 VectorStoreLoader.retrieve() - Semantic Document Search', level=2)
    doc.add_paragraph('What it does: Transforms natural language query to embedding, searches FAISS index, returns top-k relevant documents.')
    doc.add_paragraph('Why it exists: Core similarity search operation for RAG system. Separates vector search from LLM processing.')
    doc.add_paragraph('Inputs: Query string, top_k parameter (default 5).')
    doc.add_paragraph('Outputs: List of (document_text, similarity_score) tuples ordered by relevance.')
    doc.add_paragraph('Side effects: None (read-only).')
    doc.add_paragraph('What breaks if removed: No document retrieval; LLM has no context to answer queries.')
    doc.add_paragraph('Edge cases handled: (1) Invalid FAISS idx=-1: Skipped. (2) Low similarity scores: Filtered by RELEVANCE_THRESHOLD (0.3). (3) Empty index: Returns empty list.')
    doc.add_paragraph('Edge cases NOT handled: (1) Embedding model not loaded: NoneType error. (2) Metadata keys mismatch: KeyError.')
    
    doc.add_heading('5.3 LLMProcessor.process_query() - LLM with Guardrails', level=2)
    doc.add_paragraph('What it does: Applies multiple guardrails to ensure LLM only uses provided context, never hallucinates.')
    doc.add_paragraph('Why it exists: Prevents model from generating fabricated information outside knowledge base.')
    doc.add_paragraph('Inputs: query, retrieved_documents, profile_id, optional conversation_context.')
    doc.add_paragraph('Outputs: Natural language string answer.')
    doc.add_paragraph('Guardrails: (1) Context check - no context returns fallback. (2) Context string - formats docs with attribution. (3) Prompt isolation - marks system rules. (4) Not-found detection - replaces LLM "not found" with standard fallback.')
    doc.add_paragraph('Side effects: Makes HTTP POST to Ollama. May timeout (500s limit). May consume significant memory for large documents.')
    doc.add_paragraph('What breaks if removed: LLM becomes unrestricted and may hallucinate / violate client data isolation.')
    
    doc.add_page_break()
    
    # SECTION 6: Functions, Classes & Methods
    doc.add_heading('SECTION 6: Functions, Classes & Methods', level=1)
    
    doc.add_heading('6.1 QueryService Class', level=2)
    doc.add_paragraph('Purpose: Orchestrate query processing pipeline.')
    doc.add_paragraph('Methods:')
    doc.add_paragraph('- process(payload: dict) -> dict: Main entry point. Parameters: payload with client_id, query, conversation_context. Returns: {status, answer, context_retrieved}. Error handling: ValueError from validator propagates upstream; LLM errors contained and returned. Dependencies: ProfileManager, LLMProcessor, success_response utility.')
    
    doc.add_heading('6.2 VectorStoreLoader Class', level=2)
    doc.add_paragraph('Purpose: Load and manage FAISS index for a profile.')
    doc.add_paragraph('Constructor: __init__(profile_id: str). Validates profile ID against PROFILES dict. Loads embedding model and FAISS index. Raises ValueError if profile invalid, FileNotFoundError if index missing.')
    doc.add_paragraph('Methods:')
    doc.add_paragraph('- _load_embedder(): Loads sentence-transformer model. Side effect: Downloads model if not cached. Returns: None.')
    doc.add_paragraph('- _load_vector_store(): Loads FAISS index and metadata pickle. Raises: FileNotFoundError if files missing.')
    doc.add_paragraph('- retrieve(query: str, top_k: int) -> List[Tuple[str, float]]: Main retrieval method. Encodes query, searches FAISS, filters by threshold. Returns: List of document-similarity tuples.')
    doc.add_paragraph('- has_documents() -> bool: Check if vector store populated. Returns: True if index.ntotal > 0.')
    
    doc.add_heading('6.3 LLMProcessor Class', level=2)
    doc.add_paragraph('Purpose: Interface to Ollama LLM with context guardrails.')
    doc.add_paragraph('Constructor: __init__(base_url: str). Sets Ollama endpoint. Stores model name and API endpoint.')
    doc.add_paragraph('Key Methods:')
    doc.add_paragraph('- process_query(query, retrieved_documents, profile_id, conversation_context) -> str: Main entry point. Applies guardrails, calls Ollama, returns answer.')
    doc.add_paragraph('- _call_ollama(prompt: str, timeout: int) -> str: Makes HTTP POST to Ollama. Returns: LLM response or error message. Handles: ConnectionError, Timeout, HTTP errors.')
    doc.add_paragraph('- _build_context_string(documents) -> str: Formats documents with attribution and relevance scores.')
    doc.add_paragraph('- _build_prompt(query, context, profile_id, conversation_context) -> str: Creates final system prompt with isolation markers.')
    doc.add_paragraph('- _is_not_found_response(response: str) -> bool: Detects LLM "not found" patterns. Returns: True if response indicates information not in context.')
    
    doc.add_heading('6.4 ProfileManager Class', level=2)
    doc.add_paragraph('Purpose: Central registry for client profiles with lazy loading.')
    doc.add_paragraph('Methods:')
    doc.add_paragraph('- get_store(profile_id: str) -> VectorStoreLoader: Returns cached or newly-loaded store. Raises: ValueError if profile_id invalid.')
    doc.add_paragraph('- load_profile(profile_id: str) -> VectorStoreLoader: Alias for get_store().')
    doc.add_paragraph('- list_profiles() -> list: Returns sorted list of valid profile IDs.')
    doc.add_paragraph('- is_valid_profile(profile_id: str) -> bool: Check if profile in CLIENTS.')
    
    doc.add_heading('6.5 Config Class', level=2)
    doc.add_paragraph('Purpose: Load and expose YAML configuration as attributes.')
    doc.add_paragraph('Constructor: __init__(environment="PRODUCTION"). Loads YAML, merges configs, substitutes env vars.')
    doc.add_paragraph('Methods:')
    doc.add_paragraph('- _get_config(environment): Merge config.yml and db_config.yml for environment.')
    doc.add_paragraph('- _load_yaml(yaml_file, environment): Load YAML file, optionally filter by environment.')
    doc.add_paragraph('- _get_value(key, default): Get value with env var substitution (${VAR_NAME} pattern).')
    doc.add_paragraph('- _populate_attributes(): Convert config dict keys to class attributes via setattr().')
    
    doc.add_heading('6.6 Engine & Read_Write Classes', level=2)
    doc.add_paragraph('Engine: DatabaseTL connection management for PostgreSQL.')
    doc.add_paragraph('- connect_engine(): Creates SQLAlchemy engine and connection. Returns: (connection, 1) on success or (error_msg, -1) on failure.')
    doc.add_paragraph('- disconnect_engine(connection): Closes connection. Returns: "Success".')
    doc.add_paragraph('')
    doc.add_paragraph('Read_Write: SQL query execution.')
    doc.add_paragraph('- fetch_data(query, connection): Execute SELECT. Returns: (DataFrame, status_code 1 for success, -1 for failure).')
    doc.add_paragraph('- post_data(df, table_name, connection): INSERT DataFrame to table. Drops Id column if present. Commits transaction.')
    doc.add_paragraph('- update_data(query, connection): Execute UPDATE/DELETE. Commits transaction. Returns: (message, status_code).')
    
    doc.add_page_break()
    
    # SECTION 7-14: Remaining sections with concise content
    doc.add_heading('SECTION 7: State & Data Flow', level=1)
    doc.add_heading('7.1 Data Lifecycle', level=2)
    doc.add_paragraph('Query: (1) Submitted -> (2) Validated -> (3) Vector store loaded -> (4) Documents retrieved -> (5) LLM processes -> (6) Response returned. No persistence by default.')
    doc.add_paragraph('Session: (1) Created with unique session_id -> (2) Query results stored in PostgreSQL as JSONB -> (3) History retrieved for follow-ups -> (4) Persists until manual deletion.')
    doc.add_paragraph('File Upload: (1) Saved to data/ -> (2) Parsed, types inferred -> (3) Converted to Parquet -> (4) Metadata in PostgreSQL -> (5) Manual cleanup required.')
    doc.add_paragraph('Vector Store: (1) Pre-built offline -> (2) Lazy-loaded on first query -> (3) Cached for reuse -> (4) Lost on app restart.')
    
    doc.add_page_break()
    
    doc.add_heading('SECTION 8: APIs / Interfaces', level=1)
    doc.add_heading('8.1 REST Endpoints', level=2)
    doc.add_paragraph('POST /api/query: Query knowledge base. Request: {client_id, query, conversation_context}. Response: {status, answer, context_retrieved}.')
    doc.add_paragraph('GET /api/clients: List clients. Response: {clients: [...], total: N}.')
    doc.add_paragraph('POST /api/web-search: Web search fallback. Request: {query}. Response: {status, answer, source: "web"}.')
    doc.add_paragraph('GET /api/health: Health check. Response: {status: "healthy"}.')
    doc.add_paragraph('GET /api/create_session: Create session. Request: ?user_id=123. Response: {user_id, session_id}.')
    doc.add_paragraph('POST /api/fetch_attribute/dtype/v1.0: Upload file. Multipart: file, user_id, session_id.')
    doc.add_paragraph('POST /api/run_query: Run data analysis. Form: userid, sessionid, question, file_name.')
    
    doc.add_page_break()
    
    doc.add_heading('SECTION 9: Configuration & Constants', level=1)
    doc.add_heading('9.1 Key Configuration Values', level=2)
    doc.add_paragraph('API_PREFIX: "/api" - URL prefix for endpoints')
    doc.add_paragraph('TOP_K_DOCUMENTS: 5 - Documents to retrieve per query')
    doc.add_paragraph('RELEVANCE_THRESHOLD: 0.3 - Minimum similarity to include document')
    doc.add_paragraph('EMBEDDING_MODEL: "sentence-transformers/all-MiniLM-L6-v2" - Embedding model')
    doc.add_paragraph('LLM_MODEL: "gemma3:1b" - Ollama model')
    doc.add_paragraph('LLM_TEMPERATURE: 0.2 - Creativity (lower = factual)')
    doc.add_paragraph('OLLAMA_BASE_URL: "http://localhost:11434" - Ollama endpoint')
    doc.add_paragraph('ENABLE_NO_CONTEXT_FALLBACK: False - Allow LLM without context')
    doc.add_paragraph('ENABLE_CROSS_PROFILE_INFERENCE: False - Cross-profile queries (MUST be False)')
    
    doc.add_heading('9.2 Misconfiguration Risks', level=2)
    doc.add_paragraph('Wrong OLLAMA_BASE_URL: Connection refused on every query. Fix: Verify Ollama running.')
    doc.add_paragraph('Wrong DB credentials: Cannot connect to PostgreSQL. Fix: Test: psql -U user -d db')
    doc.add_paragraph('Missing FAISS index: FileNotFoundError on profile load. Fix: Verify files in src/vector_stores/')
    doc.add_paragraph('Embedding model not downloaded: Timeout on first query. Fix: Pre-download model.')
    doc.add_paragraph('ENABLE_CROSS_PROFILE_INFERENCE = True: Data leakage between clients. Fix: Always False.')
    
    doc.add_page_break()
    
    doc.add_heading('SECTION 10: Error Handling & Logging', level=1)
    doc.add_heading('10.1 Handled Errors', level=2)
    doc.add_paragraph('ValueError (invalid client_id): @app.errorhandler -> 400 response')
    doc.add_paragraph('FileNotFoundError (FAISS missing): Raised during load; caught upstream -> 500')
    doc.add_paragraph('ConnectionError (Ollama unreachable): Caught in _call_ollama() -> error message')
    doc.add_paragraph('Timeout (Ollama slow): Caught; returns timeout error and guidance')
    doc.add_paragraph('psycopg2.OperationalError: Caught with -1 status code; recovery needed')
    
    doc.add_heading('10.2 Current Logging', level=2)
    doc.add_paragraph('Method: print() statements throughout codebase.')
    doc.add_paragraph('Issues: (1) No level filtering. (2) No timestamps. (3) No log aggregation. (4) Mixed stdout/stderr.')
    doc.add_paragraph('Recommendation: Migrate to logging module with configuration.')
    
    doc.add_page_break()
    
    doc.add_heading('SECTION 11: Security Considerations', level=1)
    doc.add_heading('11.1 Data Isolation', level=2)
    doc.add_paragraph('Client Isolation: Each client_id has isolated knowledge base. QueryValidator enforces valid client_id.')
    doc.add_paragraph('Vulnerability: ENABLE_CROSS_PROFILE_INFERENCE = True breaks isolation.')
    doc.add_paragraph('Mitigation: Enforce as False; code review.')
    
    doc.add_heading('11.2 Secrets Handling', level=2)
    doc.add_paragraph('Database Password: In plaintext in db_config.yml. Risk: File access exposes password.')
    doc.add_paragraph('Mitigation: Use environment variables: ${DB_PASSWORD}. Use secrets manager. chmod 600 config files.')
    doc.add_paragraph('Google API Keys: Hardcoded in tatvaAi_utils.py. Risk: Public in repo; quota exceeded.')
    doc.add_paragraph('Mitigation: Use environment variables; secrets management.')
    
    doc.add_heading('11.3 Known Vulnerabilities', level=2)
    doc.add_paragraph('Debug Mode Enabled: Flask runs with debug=True. Stack traces exposed. Fix: debug=False in production.')
    doc.add_paragraph('No Rate Limiting: Single client can spam requests. Fix: Add flask-limiter.')
    doc.add_paragraph('SQL Injection Risk: User provides column names. Fix: LLM generates SQL (typically safe) but validate.')
    
    doc.add_page_break()
    
    doc.add_heading('SECTION 12: Performance & Scalability', level=1)
    doc.add_heading('12.1 Performance Characteristics', level=2)
    doc.add_paragraph('Query embedding: 10-50ms')
    doc.add_paragraph('FAISS search: 1-5ms')
    doc.add_paragraph('Ollama inference: 2-30s (PRIMARY BOTTLENECK)')
    doc.add_paragraph('PostgreSQL query: 10-100ms')
    doc.add_paragraph('Full pipeline: 3-35s')
    
    doc.add_heading('12.2 Scaling Limits', level=2)
    doc.add_paragraph('Single-server: ~100 concurrent requests with 8 threads.')
    doc.add_paragraph('FAISS: 10M docs = ~10GB RAM. Scaling above requires distributed search.')
    doc.add_paragraph('PostgreSQL: ~1000 TPS single instance. Scaling requires replication.')
    doc.add_paragraph('Ollama: Single instance serves serially. Multiple instances need load balancer.')
    
    doc.add_heading('12.3 Recommendations', level=2)
    doc.add_paragraph('1. Pre-load embedding model at startup.')
    doc.add_paragraph('2. Use Redis for caching frequent queries.')
    doc.add_paragraph('3. Parallelize Ollama via pools or async.')
    doc.add_paragraph('4. Use database indexes on frequently queried columns.')
    doc.add_paragraph('5. Profile code with cProfile.')
    
    doc.add_page_break()
    
    doc.add_heading('SECTION 13: Limitations & Technical Debt', level=1)
    doc.add_heading('13.1 Known Weaknesses', level=2)
    doc.add_paragraph('No authentication: Anyone with API access can query any profile. Fix: Add API key auth / JWT.')
    doc.add_paragraph('FAISS not updateable: Knowledge bases become stale. Fix: Use PostgreSQL vector extension.')
    doc.add_paragraph('String conversation context: Lose nuance in follow-ups. Fix: Structured session objects.')
    doc.add_paragraph('No observability: Cannot monitor production. Fix: Add Prometheus; OpenTelemetry.')
    doc.add_paragraph('Hardcoded paths: Not portable. Fix: Environment-specific config.')
    
    doc.add_heading('13.2 Refactor Opportunities', level=2)
    doc.add_paragraph('Extract file handling into FileService class.')
    doc.add_paragraph('Replace print() with logging module.')
    doc.add_paragraph('Move database decorators to context manager.')
    doc.add_paragraph('Standardize error handling (exception hierarchy).')
    doc.add_paragraph('Extract LLM prompting into separate PromptBuilder.')
    doc.add_paragraph('Move API keys to environment variables.')
    
    doc.add_page_break()
    
    doc.add_heading('SECTION 14: Extension & Modification Guide', level=1)
    
    doc.add_heading('14.1 Adding New Features', level=2)
    doc.add_paragraph('Pattern: (1) Create controller blueprint. (2) Create service class. (3) Add validator if needed. (4) Register blueprint with url_prefix.')
    doc.add_paragraph('Example: Add analytics endpoint: Create analytics_controller.py -> analytics_service.py -> register in app.py.')
    
    doc.add_heading('14.2 Adding New Client Profile', level=2)
    doc.add_paragraph('Steps: (1) Build FAISS index offline. (2) Place files: client_id_index.faiss, client_id_metadata.pkl. (3) Add to CLIENTS dict in settings.py. (4) Test via /api/clients and /api/query.')
    
    doc.add_heading('14.3 Modifying LLM Behavior', level=2)
    doc.add_paragraph('Change Model: LLM_MODEL in settings.py. E.g., "llama2:13b".')
    doc.add_paragraph('Change Temperature: LLM_TEMPERATURE in settings.py (0.2=factual, 0.9=creative).')
    doc.add_paragraph('Change System Prompt: Edit SYSTEM_CONTRACT in processor.py.')
    
    doc.add_heading('14.4 Adding Authentication', level=2)
    doc.add_paragraph('Steps: (1) Add @require_auth decorator. (2) Extract JWT from headers. (3) Validate token. (4) Enforce access control by client.')
    
    doc.add_heading('14.5 Adding Caching', level=2)
    doc.add_paragraph('Pattern: Check Redis before calling LLM. Store result with TTL. Return cached on hit.')
    
    doc.add_heading('14.6 Patterns to Follow', level=2)
    doc.add_paragraph('- Separation of Concerns: Controllers -> Services -> Utilities.')
    doc.add_paragraph('- Static Service Methods: Enable easy import and testing.')
    doc.add_paragraph('- Lazy Loading: Reduce startup time.')
    doc.add_paragraph('- Configuration as Source of Truth: settings.py centralized.')
    
    doc.add_heading('14.7 Anti-patterns to Avoid', level=2)
    doc.add_paragraph('- Global mutable state (not thread-safe).')
    doc.add_paragraph('- Mixing HTTP and business logic in controllers.')
    doc.add_paragraph('- Hardcoding paths or credentials.')
    doc.add_paragraph('- Catching generic Exception.')
    doc.add_paragraph('- Synchronous blocking calls (use async/await or pools).')
    
    doc.add_page_break()
    
    # Conclusion
    doc.add_heading('CONCLUSION', level=1)
    doc.add_paragraph(
        'PlamBO Backend is a sophisticated multi-tenant knowledge management and data analysis system combining '
        'vector search, large language models, and relational databases. The architecture emphasizes client isolation, '
        'semantic search accuracy, and guardrails to prevent hallucination.'
    )
    doc.add_paragraph(
        'Key Takeaways:'
    )
    doc.add_paragraph('1. All endpoints validate client_id for data isolation.', style='List Number')
    doc.add_paragraph('2. Vector stores are lazy-loaded for efficiency.', style='List Number')
    doc.add_paragraph('3. Ollama provides local LLM inference without external dependencies.', style='List Number')
    doc.add_paragraph('4. PostgreSQL persists sessions and metadata.', style='List Number')
    doc.add_paragraph('5. Extensive configuration enables tuning.', style='List Number')
    doc.add_paragraph('6. Security and performance optimizations are ongoing opportunities.', style='List Number')
    
    doc.add_paragraph()
    doc.add_paragraph('For deployment: Ensure PostgreSQL, Ollama, and FAISS indexes are ready. Configure settings.py with profiles. Pre-download models. Test end-to-end. Monitor logs and error rates. Plan scaling via async, caching, and distributed infrastructure.')
    
    # Save document
    output_path = 'PlamBO_Backend_Technical_Documentation.docx'
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    try:
        output = create_documentation()
        print(f'✓ Documentation successfully created: {output}')
    except Exception as e:
        print(f'✗ Error creating documentation: {str(e)}')
        import traceback
        traceback.print_exc()
